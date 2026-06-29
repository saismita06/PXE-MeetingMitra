"""
Document processing and conversion services.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor



def process_markdown_to_docx(doc, content):
    """Convert markdown content to properly formatted Word document elements.

    Supports:
    - Tables (markdown pipe tables)
    - Headings (# ## ###)
    - Bold text (**text**)
    - Italic text (*text* or _text_)
    - Bold italic (***text***)
    - Inline code (`code`)
    - Code blocks (```code```)
    - Strikethrough (~~text~~)
    - Links ([text](url))
    - Bullet lists (- or *)
    - Numbered lists (1. 2. 3.)
    - Horizontal rules (--- or ***)
    """
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    import re

    def ensure_unicode_font(run, text):
        """Ensure the run uses a font that supports the characters in the text."""
        # Check if text contains non-ASCII characters
        try:
            text.encode('ascii')
            # Text is pure ASCII, no special font needed
        except UnicodeEncodeError:
            # Text contains non-ASCII characters, use a font with better Unicode support
            # Use Arial for broad compatibility - it has good Unicode support on most systems
            run.font.name = 'Arial'
            # Set the East Asian font for CJK (Chinese, Japanese, Korean) text
            # This ensures proper rendering in Word
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        return run

    def add_formatted_run(paragraph, text):
        """Add a run with inline formatting to a paragraph."""
        if not text:
            return

        # Pattern for all inline formatting
        # Order matters: check triple asterisk before double/single
        patterns = [
            (r'\*\*\*(.*?)\*\*\*', lambda p, t: (lambda r: (setattr(r, 'bold', True), setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Bold italic
            (r'\*\*(.*?)\*\*', lambda p, t: (lambda r: (setattr(r, 'bold', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Bold
            (r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)', lambda p, t: (lambda r: (setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Italic with *
            (r'\b_(.*?)_\b', lambda p, t: (lambda r: (setattr(r, 'italic', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Italic with _
            (r'~~(.*?)~~', lambda p, t: (lambda r: (setattr(r, 'strike', True), ensure_unicode_font(r, t)))(p.add_run(t))),  # Strikethrough
            (r'`([^`]+)`', lambda p, t: add_code_run(p, t)),  # Inline code
            (r'\[([^\]]+)\]\(([^)]+)\)', lambda p, t, u: add_link_run(p, t, u)),  # Links
        ]

        def add_code_run(para, text):
            """Add inline code with monospace font and background."""
            run = para.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson color for code
            # Check if we need Unicode support for code
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                # Use Consolas as fallback for better Unicode support in monospace
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            return run

        def add_link_run(para, text, url):
            """Add a hyperlink-styled run (note: actual hyperlinks require more complex handling)."""
            full_text = f"{text} ({url})"
            run = para.add_run(full_text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
            run.font.underline = True
            ensure_unicode_font(run, full_text)
            return run

        # Process the text with all patterns
        remaining_text = text
        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            matched_pattern = None

            # Find the earliest matching pattern
            for pattern, handler in patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    matched_pattern = handler

            if earliest_match:
                # Add text before the match
                if earliest_pos > 0:
                    run = paragraph.add_run(remaining_text[:earliest_pos])
                    ensure_unicode_font(run, remaining_text[:earliest_pos])

                # Apply formatting for the matched text
                if '[' in earliest_match.group(0) and '](' in earliest_match.group(0):
                    # Special handling for links (two groups)
                    matched_pattern(paragraph, earliest_match.group(1), earliest_match.group(2))
                else:
                    matched_pattern(paragraph, earliest_match.group(1))

                # Continue with remaining text
                remaining_text = remaining_text[earliest_match.end():]
            else:
                # No more patterns, add the rest as plain text
                run = paragraph.add_run(remaining_text)
                ensure_unicode_font(run, remaining_text)
                break

    def parse_table(lines, start_idx):
        """Parse a markdown table starting at the given index."""
        if start_idx >= len(lines):
            return None, start_idx

        # Check if this looks like a table
        if '|' not in lines[start_idx]:
            return None, start_idx

        table_data = []
        idx = start_idx

        while idx < len(lines) and '|' in lines[idx]:
            # Skip separator lines
            if re.match(r'^[\s\|\-:]+$', lines[idx]):
                idx += 1
                continue

            # Parse cells
            cells = [cell.strip() for cell in lines[idx].split('|')]
            # Remove empty cells at start and end
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]

            if cells:
                table_data.append(cells)
            idx += 1

        if table_data:
            return table_data, idx
        return None, start_idx

    # Split content into lines
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                # End of code block - add it as preformatted text
                in_code_block = False
                if code_block_content:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    code_text = '\n'.join(code_block_content)
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)
                    # Check if we need Unicode support for code blocks
                    try:
                        code_text.encode('ascii')
                    except UnicodeEncodeError:
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Check for table
        table_data, end_idx = parse_table(lines, i)
        if table_data:
            # Create Word table
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'

            # Populate table
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        # Clear existing paragraphs and add new one
                        cell.text = ""
                        p = cell.add_paragraph()
                        add_formatted_run(p, cell_text)
                        # Make header row bold
                        if row_idx == 0:
                            for run in p.runs:
                                run.bold = True

            doc.add_paragraph('')  # Space after table
            i = end_idx
            continue

        line = line.rstrip()

        # Skip empty lines
        if not line:
            doc.add_paragraph('')
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^(\*{3,}|-{3,}|_{3,})$', line.strip()):
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        # Bullet points
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            # Get the indentation level
            indent = len(line) - len(line.lstrip())
            bullet_text = line.lstrip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Add indentation if nested
            if indent > 0:
                p.paragraph_format.left_indent = Pt(indent * 10)
            add_formatted_run(p, bullet_text)
        # Numbered lists
        elif re.match(r'^\s*\d+\.', line):
            match = re.match(r'^(\s*)(\d+)\.\s*(.*)', line)
            if match:
                indent = len(match.group(1))
                list_text = match.group(3)
                p = doc.add_paragraph(style='List Number')
                if indent > 0:
                    p.paragraph_format.left_indent = Pt(indent * 10)
                add_formatted_run(p, list_text)
        # Blockquote
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(30)
            add_formatted_run(p, line[2:])
            # Add a gray color to indicate quote
            for run in p.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            add_formatted_run(p, line)

        i += 1

# --- Database Models ---
# --- Database Models ---
# Models have been extracted to src/models/ and imported at the top of this file

# --- Forms for Authentication ---
# --- Custom Password Validator ---
# password_check utility has been extracted to src/utils/security.py


# --- Blueprint Registration ---
# Import and register all blueprints for modular route organization



