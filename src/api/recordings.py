"""
Recording upload, processing, and management.

This blueprint was auto-generated from app.py route extraction.
"""

import os
import json
import re
import mimetypes
import time
import threading
import subprocess
from datetime import datetime, timedelta
from src.services.job_queue import job_queue
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, Response, current_app, make_response
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from sqlalchemy import select
from email.utils import encode_rfc2231

from src.database import db
from src.models import *
from src.utils import *
from src.config.app_config import ASR_MIN_SPEAKERS, ASR_MAX_SPEAKERS, ASR_DIARIZE, USE_NEW_TRANSCRIPTION_ARCHITECTURE


def _resolve_transcription_model(value):
    """Apply admin-curated allowlist + default to a candidate model name.

    Resolution:
      1. If the caller passed a non-empty value:
         a. If the admin saved a visible-models list in system_setting and the
            value is in it, accept.
         b. Else if TRANSCRIPTION_MODELS_AVAILABLE is set and the value is in
            it, accept.
         c. Else if neither list is configured, accept (no allowlist enforced).
         d. Otherwise, drop the value with a warning.
      2. If the caller passed nothing, fall back to the admin-saved default
         model (system_setting key `transcription_default_model`) when set.
    Returns the validated model id or None.
    """
    from src.config.app_config import TRANSCRIPTION_MODELS_AVAILABLE
    import json as _json

    candidate = (value or '').strip() or None
    visible = []
    try:
        raw = SystemSetting.get_setting('transcription_models_visible_json', None)
        if raw:
            parsed = _json.loads(raw)
            if isinstance(parsed, list):
                visible = [
                    (item['value'] if isinstance(item, dict) else item)
                    for item in parsed if item
                ]
    except Exception:
        visible = []

    if candidate:
        in_db_list = bool(visible) and candidate in visible
        in_env_list = bool(TRANSCRIPTION_MODELS_AVAILABLE) and candidate in TRANSCRIPTION_MODELS_AVAILABLE
        if visible or TRANSCRIPTION_MODELS_AVAILABLE:
            if not (in_db_list or in_env_list):
                current_app.logger.warning(
                    f"Ignoring transcription_model={candidate!r} — not in admin-curated list or TRANSCRIPTION_MODELS_AVAILABLE"
                )
                candidate = None
        return candidate

    # Fall back to admin-saved default when no override flowed through.
    default = SystemSetting.get_setting('transcription_default_model', None)
    return default or None
from src.tasks.processing import format_transcription_for_llm
from src.utils.ffmpeg_utils import FFmpegError, FFmpegNotFoundError
from src.utils.titles import resolve_upload_title
from src.services.speaker import update_speaker_usage, identify_unidentified_speakers_from_text
from src.services.speaker_embedding_matcher import update_speaker_embedding
from src.services.speaker_snippets import create_speaker_snippets

# Incognito mode - disabled by default, enable via environment variable
ENABLE_INCOGNITO_MODE = os.environ.get('ENABLE_INCOGNITO_MODE', 'false').lower() == 'true'
from src.services.document import process_markdown_to_docx
from src.services.llm import client, chat_client, call_llm_completion, call_chat_completion, process_streaming_with_thinking, TokenBudgetExceeded
from src.services.embeddings import process_recording_chunks
from src.file_exporter import export_recording, mark_export_as_deleted
from src.utils.ffprobe import get_codec_info, get_creation_date, get_duration, FFProbeError
from src.utils.audio_conversion import convert_if_needed
from src.services.storage import get_storage_service
from src.utils.file_hash import compute_file_sha256

# Create blueprint
recordings_bp = Blueprint('recordings', __name__)

# Configuration from environment
ENABLE_INQUIRE_MODE = os.environ.get('ENABLE_INQUIRE_MODE', 'false').lower() == 'true'
ENABLE_AUTO_DELETION = os.environ.get('ENABLE_AUTO_DELETION', 'false').lower() == 'true'
DELETION_MODE = os.environ.get('DELETION_MODE', 'full_recording')  # 'audio_only' or 'full_recording'
USERS_CAN_DELETE = os.environ.get('USERS_CAN_DELETE', 'true').lower() == 'true'
ENABLE_INTERNAL_SHARING = os.environ.get('ENABLE_INTERNAL_SHARING', 'false').lower() == 'true'
VIDEO_RETENTION = os.environ.get('VIDEO_RETENTION', 'false').lower() == 'true'
VIDEO_PASSTHROUGH_ASR = os.environ.get('VIDEO_PASSTHROUGH_ASR', 'false').lower() == 'true'
USE_ASR_ENDPOINT = os.environ.get('USE_ASR_ENDPOINT', 'false').lower() == 'true'
ENABLE_CHUNKING = os.environ.get('ENABLE_CHUNKING', 'true').lower() == 'true'

def reindex_recording_chunks_async(recording_id):
    """Rebuild a recording's semantic-search (Inquire) chunks in the background.

    The Inquire/RAG chunks are a snapshot of recording.transcription. When the
    transcription text changes — speaker names applied, transcript edited — the
    chunks must be rebuilt or Inquire keeps answering from the stale text (e.g.
    'SPEAKER_00' instead of the applied name). Runs in a daemon thread so the
    edit's HTTP response isn't blocked on embedding generation; no-op when
    Inquire mode is off. process_recording_chunks deletes-then-recreates with
    rollback safety, so a failure leaves the previous chunks intact.
    """
    if not ENABLE_INQUIRE_MODE:
        return
    app_obj = current_app._get_current_object()

    def _run():
        with app_obj.app_context():
            try:
                process_recording_chunks(recording_id)
                app_obj.logger.info(f"Reindexed Inquire chunks for recording {recording_id} after edit")
            except Exception as e:
                app_obj.logger.error(f"Failed to reindex Inquire chunks for recording {recording_id}: {e}", exc_info=True)

    threading.Thread(target=_run, name=f"reindex-chunks-{recording_id}", daemon=True).start()


# Global helpers (will be injected from app)
has_recording_access = None
get_user_recording_status = None
set_user_recording_status = None
enrich_recording_dict_with_user_status = None
bcrypt = None
csrf = None
limiter = None
chunking_service = None

def init_recordings_helpers(**kwargs):
    """Initialize helper functions and extensions from app."""
    global has_recording_access, get_user_recording_status, set_user_recording_status, enrich_recording_dict_with_user_status, bcrypt, csrf, limiter, chunking_service
    has_recording_access = kwargs.get('has_recording_access')
    get_user_recording_status = kwargs.get('get_user_recording_status')
    set_user_recording_status = kwargs.get('set_user_recording_status')
    enrich_recording_dict_with_user_status = kwargs.get('enrich_recording_dict_with_user_status')
    bcrypt = kwargs.get('bcrypt')
    csrf = kwargs.get('csrf')
    limiter = kwargs.get('limiter')
    chunking_service = kwargs.get('chunking_service')


# --- Routes ---

@recordings_bp.route('/recording/<int:recording_id>/download/transcript')
@login_required
def download_transcript_with_template(recording_id):
    """Download transcript with custom template formatting."""
    try:
        import re
        from datetime import timedelta

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for this recording'}), 400

        # Get template ID from query params
        template_id = request.args.get('template_id', type=int)

        # Get the template
        if template_id:
            template = TranscriptTemplate.query.filter_by(
                id=template_id,
                user_id=current_user.id
            ).first()
        else:
            # Use default template
            template = TranscriptTemplate.query.filter_by(
                user_id=current_user.id,
                is_default=True
            ).first()

        # If no template found, use a basic format
        if not template:
            template_format = "[{{speaker}}]: {{text}}"
        else:
            template_format = template.template

        # Helper functions for formatting
        def format_time(seconds):
            """Format seconds to HH:MM:SS"""
            if seconds is None:
                return "00:00:00"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"

        def format_srt_time(seconds):
            """Format seconds to SRT format HH:MM:SS,mmm"""
            if seconds is None:
                return "00:00:00,000"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            millis = int((td.total_seconds() % 1) * 1000)
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

        # Parse transcription - handle both JSON (diarized) and plain text formats
        is_diarized = False
        transcription_data = None
        try:
            transcription_data = json.loads(recording.transcription)
            if isinstance(transcription_data, list):
                is_diarized = True
        except (json.JSONDecodeError, TypeError):
            # Not JSON, treat as plain text
            pass

        # If plain text transcription, return it as-is (no template formatting applies)
        if not is_diarized:
            formatted_transcript = recording.transcription
        else:
            # Generate formatted transcript from diarized segments
            output_lines = []
            for index, segment in enumerate(transcription_data, 1):
                line = template_format

                # Replace variables
                replacements = {
                    '{{index}}': str(index),
                    '{{speaker}}': segment.get('speaker', 'Unknown'),
                    '{{text}}': segment.get('sentence', ''),
                    '{{start_time}}': format_time(segment.get('start_time')),
                    '{{end_time}}': format_time(segment.get('end_time')),
                }

                for key, value in replacements.items():
                    line = line.replace(key, value)

                # Handle filters
                # Upper case filter
                line = re.sub(r'{{(.*?)\|upper}}', lambda m: replacements.get('{{' + m.group(1) + '}}', '').upper(), line)
                # SRT time filter
                line = re.sub(r'{{start_time\|srt}}', format_srt_time(segment.get('start_time')), line)
                line = re.sub(r'{{end_time\|srt}}', format_srt_time(segment.get('end_time')), line)

                output_lines.append(line)

            # Join lines
            formatted_transcript = '\n'.join(output_lines)

        # Create response
        response = make_response(formatted_transcript)
        if is_diarized and template:
            filename = f"{recording.title or 'transcript'}_{template.name}.txt"
        elif is_diarized:
            filename = f"{recording.title or 'transcript'}_formatted.txt"
        else:
            # Plain text transcription
            filename = f"{recording.title or 'transcript'}.txt"
        filename = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        current_app.logger.error(f"Error downloading transcript: {e}")
        return jsonify({'error': 'Failed to generate transcript download'}), 500




@recordings_bp.route('/recording/<int:recording_id>/download/summary')
@login_required
def download_summary_word(recording_id):
    """Download recording summary as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.summary:
            return jsonify({'error': 'No summary available for this recording'}), 400

        # Create Word document
        doc = Document()

        # Add title
        title_text = f'Summary: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            # Title contains non-ASCII characters
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        visible_tags = recording.get_visible_tags(current_user)
        if visible_tags:
            tags_str = ', '.join([tag.name for tag in visible_tags])
            add_unicode_paragraph(doc, f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line

        # Process markdown content using the helper function
        process_markdown_to_docx(doc, recording.summary)

        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'summary-{safe_title}.docx' if safe_title else f'summary-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['summary-.docx', 'summary-recording-.docx']:
            ascii_filename = f'summary-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response

    except Exception as e:
        current_app.logger.error(f"Error generating summary Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/download/chat', methods=['POST'])
@login_required
def download_chat_word(recording_id):
    """Download chat conversation as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Get chat messages from request
        data = request.json
        if not data or 'messages' not in data:
            return jsonify({'error': 'No messages provided'}), 400

        messages = data['messages']
        if not messages:
            return jsonify({'error': 'No messages to download'}), 400

        # Create Word document
        doc = Document()

        # Add title
        title_text = f'Chat Conversation: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Recording Date: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        add_unicode_paragraph(doc, f'Chat Export Date: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph('')  # Empty line

        # Add chat messages
        for message in messages:
            role = message.get('role', 'unknown')
            content = message.get('content', '')
            thinking = message.get('thinking', '')

            # Add role header
            if role == 'user':
                p = doc.add_paragraph()
                run = p.add_run('You: ')
                run.bold = True
            elif role == 'assistant':
                p = doc.add_paragraph()
                run = p.add_run('Assistant: ')
                run.bold = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f'{role.title()}: ')
                run.bold = True

            # Add thinking content if present
            if thinking and role == 'assistant':
                p = doc.add_paragraph()
                p.add_run('[Model Reasoning]\n').italic = True
                p.add_run(thinking).italic = True
                doc.add_paragraph('')  # Empty line

            # Add message content with markdown formatting
            process_markdown_to_docx(doc, content)

            doc.add_paragraph('')  # Empty line between messages

        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'chat-{safe_title}.docx' if safe_title else f'chat-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['chat-.docx', 'chat-recording-.docx']:
            ascii_filename = f'chat-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response

    except Exception as e:
        current_app.logger.error(f"Error generating chat Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/download/notes')
@login_required
def download_notes_word(recording_id):
    """Download recording notes as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.notes:
            return jsonify({'error': 'No notes available for this recording'}), 400

        # Create Word document
        doc = Document()

        # Add title
        title_text = f'Notes: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        visible_tags = recording.get_visible_tags(current_user)
        if visible_tags:
            tags_str = ', '.join([tag.name for tag in visible_tags])
            add_unicode_paragraph(doc, f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line

        # Process markdown content using the helper function
        process_markdown_to_docx(doc, recording.notes)

        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'notes-{safe_title}.docx' if safe_title else f'notes-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['notes-.docx', 'notes-recording-.docx']:
            ascii_filename = f'notes-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response

    except Exception as e:
        current_app.logger.error(f"Error generating notes Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/generate_summary', methods=['POST'])
@login_required
def generate_summary_endpoint(recording_id):
    """Generate summary for a recording that doesn't have one.

    Optional JSON body:
      - custom_prompt (string): user-supplied summarization instructions.
      - prompt_mode ('replace' | 'append', default 'replace'): in 'replace'
        mode the custom prompt overrides the resolved default; in 'append'
        mode it is appended after the resolved default as additional context.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to generate summary for this recording'}), 403

        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400

        # Check if transcription is an error message (not actual content)
        if is_transcription_error(recording.transcription):
            return jsonify({'error': 'Cannot generate summary: transcription failed. Please reprocess the transcription first.'}), 400

        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400

        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503

        # Optional custom prompt + mode (issue / discussion #253)
        data = request.get_json(silent=True) or {}
        raw_prompt = data.get('custom_prompt')
        custom_prompt = raw_prompt.strip() if isinstance(raw_prompt, str) and raw_prompt.strip() else None
        prompt_mode = (data.get('prompt_mode') or 'replace').strip().lower()
        if prompt_mode not in ('replace', 'append'):
            prompt_mode = 'replace'
        custom_prompt_append = bool(custom_prompt) and prompt_mode == 'append'

        current_app.logger.info(
            f"Queueing summary generation for recording {recording_id}"
            + (f" with custom prompt (mode={prompt_mode}, length={len(custom_prompt)})" if custom_prompt else "")
        )

        # Queue summary generation job
        job_params = {'user_id': current_user.id}
        if custom_prompt:
            job_params['custom_prompt'] = custom_prompt
            job_params['custom_prompt_append'] = custom_prompt_append

        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='summarize',
            params=job_params,
        )

        return jsonify({
            'success': True,
            'message': 'Summary generation queued'
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error starting summary generation for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/update_speakers', methods=['POST'])
@login_required
def update_speakers(recording_id):
    """Updates speaker labels in a transcription with provided names."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        speaker_map = data.get('speaker_map')
        regenerate_summary = data.get('regenerate_summary', False)

        if speaker_map is None:
            return jsonify({'error': 'No speaker map provided'}), 400

        transcription_text = recording.transcription
        is_json = False
        try:
            transcription_data = json.loads(transcription_text)
            # Updated check for our new simplified JSON format (a list of segment objects)
            is_json = isinstance(transcription_data, list)
        except (json.JSONDecodeError, TypeError):
            is_json = False

        speaker_names_used = []

        if is_json:
            # Handle new simplified JSON transcript (list of segments)
            for segment in transcription_data:
                original_speaker_label = segment.get('speaker')
                if original_speaker_label in speaker_map:
                    new_name_info = speaker_map[original_speaker_label]
                    new_name = new_name_info.get('name', '').strip()
                    # If isMe is checked but no name provided, use current user's name
                    if new_name_info.get('isMe') and not new_name:
                        new_name = current_user.name or 'Me'

                    if new_name:
                        segment['speaker'] = new_name
                        if new_name not in speaker_names_used:
                            speaker_names_used.append(new_name)

            recording.transcription = json.dumps(transcription_data)

            # Update participants only from speakers that were actually given names (not default labels)
            final_speakers = set()
            for seg in transcription_data:
                speaker = seg.get('speaker')
                if speaker and str(speaker).strip():
                    # Only include speakers that have been given actual names (not default labels like "SPEAKER_01", "SPEAKER_09", etc.)
                    # Check if this speaker was updated with a real name (not a default label)
                    if not re.match(r'^SPEAKER_\d+$', str(speaker), re.IGNORECASE):
                        final_speakers.add(speaker)
            recording.participants = ', '.join(sorted(list(final_speakers)))

        else:
            # Handle plain text transcript
            new_participants = []
            for speaker_label, new_name_info in speaker_map.items():
                new_name = new_name_info.get('name', '').strip()
                # If isMe is checked but no name provided, use current user's name
                if new_name_info.get('isMe') and not new_name:
                    new_name = current_user.name or 'Me'

                if new_name:
                    transcription_text = re.sub(r'\[\s*' + re.escape(speaker_label) + r'\s*\]', f'[{new_name}]', transcription_text, flags=re.IGNORECASE)
                    if new_name not in new_participants:
                        new_participants.append(new_name)

            recording.transcription = transcription_text
            if new_participants:
                recording.participants = ', '.join(new_participants)
            speaker_names_used = new_participants

        # Update speaker usage statistics
        if speaker_names_used:
            update_speaker_usage(speaker_names_used)

        # Update speaker voice embeddings if available
        embeddings_updated = 0
        snippets_created = 0
        if recording.speaker_embeddings and speaker_map:
            try:
                # Parse embeddings from recording
                embeddings_data = json.loads(recording.speaker_embeddings) if isinstance(recording.speaker_embeddings, str) else recording.speaker_embeddings

                # Build reverse map: SPEAKER_XX -> actual name assigned
                speaker_label_to_name = {}
                for speaker_label, speaker_info in speaker_map.items():
                    name = speaker_info.get('name', '').strip()
                    # Handle isMe checkbox
                    if speaker_info.get('isMe') and not name:
                        name = current_user.name or 'Me'

                    # Only include speakers that were given real names (not SPEAKER_XX)
                    if name and not re.match(r'^SPEAKER_\d+$', name, re.IGNORECASE):
                        speaker_label_to_name[speaker_label] = name

                # Update embeddings for each identified speaker
                for speaker_label, embedding in embeddings_data.items():
                    if speaker_label in speaker_label_to_name and embedding and len(embedding) == 256:
                        speaker_name = speaker_label_to_name[speaker_label]

                        # Find or create the speaker
                        speaker = Speaker.query.filter_by(
                            user_id=current_user.id,
                            name=speaker_name
                        ).first()

                        if speaker:
                            # Update the speaker's voice embedding
                            similarity = update_speaker_embedding(speaker, embedding, recording.id)
                            embeddings_updated += 1

                            if similarity is not None:
                                current_app.logger.info(
                                    f"Updated voice profile for '{speaker_name}' "
                                    f"(similarity: {similarity*100:.1f}%)"
                                )
                            else:
                                current_app.logger.info(
                                    f"Created initial voice profile for '{speaker_name}'"
                                )

                # Create snippets for identified speakers
                if speaker_label_to_name:
                    snippets_created = create_speaker_snippets(recording.id, speaker_map)
                    if snippets_created > 0:
                        current_app.logger.info(f"Created {snippets_created} speaker snippets")

            except Exception as e:
                current_app.logger.error(f"Error updating speaker embeddings: {e}", exc_info=True)
                # Don't fail the whole request if embedding update fails

        db.session.commit()

        # Speaker names changed the transcription text — rebuild the Inquire
        # chunks so semantic search answers with the applied names, not the
        # raw SPEAKER_XX labels. Background + best-effort.
        reindex_recording_chunks_async(recording_id)

        summary_queued = False
        if regenerate_summary:
            current_app.logger.info(f"Queueing summary regeneration for recording {recording_id} after speaker update.")
            job_queue.enqueue(
                user_id=current_user.id,
                recording_id=recording.id,
                job_type='summarize',
                params={'user_id': current_user.id}
            )
            summary_queued = True

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Speakers updated successfully.',
            'recording': recording_dict,
            'summary_queued': summary_queued
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating speakers for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/update_transcript', methods=['POST'])
@login_required
def update_transcript(recording_id):
    """Updates the complete transcript data including text edits and speaker changes."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        transcript_data = data.get('transcript_data')
        speaker_map = data.get('speaker_map', {})
        regenerate_summary = data.get('regenerate_summary', False)

        if not transcript_data or not isinstance(transcript_data, list):
            return jsonify({'error': 'Invalid transcript data provided'}), 400

        # Update speaker names in the transcript data
        speaker_names_used = []
        for segment in transcript_data:
            original_speaker_label = segment.get('speaker')

            # Apply speaker name mapping if provided
            if original_speaker_label in speaker_map:
                new_name_info = speaker_map[original_speaker_label]
                new_name = new_name_info.get('name', '').strip()
                if new_name_info.get('isMe'):
                    new_name = current_user.name or 'Me'

                if new_name:
                    segment['speaker'] = new_name
                    if new_name not in speaker_names_used:
                        speaker_names_used.append(new_name)

        # Save the updated transcript
        recording.transcription = json.dumps(transcript_data)

        # Update participants
        final_speakers = set()
        for seg in transcript_data:
            speaker = seg.get('speaker')
            if speaker and str(speaker).strip():
                # Only include speakers with real names (not default labels)
                if not re.match(r'^SPEAKER_\d+$', str(speaker), re.IGNORECASE):
                    final_speakers.add(speaker)
        recording.participants = ', '.join(sorted(list(final_speakers)))

        # Update speaker usage statistics
        if speaker_names_used:
            update_speaker_usage(speaker_names_used)

        db.session.commit()

        # The transcript text changed — rebuild Inquire chunks so semantic
        # search reflects the edits (names, corrections). Background + best-effort.
        reindex_recording_chunks_async(recording_id)

        summary_queued = False
        if regenerate_summary:
            current_app.logger.info(f"Queueing summary regeneration for recording {recording_id} after transcript update.")
            job_queue.enqueue(
                user_id=current_user.id,
                recording_id=recording.id,
                job_type='summarize',
                params={'user_id': current_user.id}
            )
            summary_queued = True
            # Export will happen after summary regenerates
        else:
            # Re-export the recording if auto-export is enabled
            export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Transcript updated successfully.',
            'recording': recording_dict,
            'summary_queued': summary_queued
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating transcript for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/auto_identify_speakers', methods=['POST'])
@login_required
def auto_identify_speakers(recording_id):
    """
    Automatically identifies speakers in a transcription using an LLM.
    Strips existing names and re-identifies all speakers from scratch.
    """
    from src.services.speaker_identification import identify_speakers_from_transcript

    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for speaker identification'}), 400

        try:
            transcription_data = json.loads(recording.transcription)
        except (json.JSONDecodeError, TypeError):
            return jsonify({'error': 'Transcription format not supported for auto-identification'}), 400

        if not isinstance(transcription_data, list):
            return jsonify({'error': 'Transcription format not supported for auto-identification'}), 400

        speaker_map = identify_speakers_from_transcript(transcription_data, current_user.id)

        if not speaker_map:
            return jsonify({'error': 'No speakers found in transcription'}), 400

        return jsonify({'success': True, 'speaker_map': speaker_map})

    except ValueError as ve:
        return jsonify({'error': str(ve)}), 503
    except Exception as e:
        current_app.logger.error(f"Error during auto speaker identification for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': f'An unexpected error occurred: {str(e)}'}), 500

# --- Chat with Transcription ---


@recordings_bp.route('/recording/<int:recording_id>/reprocess_transcription', methods=['POST'])
@login_required
def reprocess_transcription(recording_id):
    """Reprocess transcription for a given recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403

        if not recording.audio_path or not get_storage_service().exists(recording.audio_path):
            return jsonify({'error': 'Audio file not found for reprocessing'}), 404

        if recording.status in ['QUEUED', 'PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400

        # File path and name for processing (conversion handled in background task if needed)
        filepath = recording.audio_path
        filename_for_asr = recording.original_filename or os.path.basename(filepath)

        # --- Proceed with reprocessing ---
        recording.transcription = None
        recording.summary = None
        recording.status = 'QUEUED'  # Will change to PROCESSING when job starts

        # Clear existing events since they depend on the transcription
        Event.query.filter_by(recording_id=recording_id).delete()

        db.session.commit()

        current_app.logger.info(f"Queueing transcription reprocessing for recording {recording_id}")

        # Prepare job parameters
        data = request.json or {}
        start_time = datetime.utcnow()
        app_context = current_app._get_current_object().app_context()

        # Build job parameters - language handling:
        # - If 'language' key exists with a value (e.g., 'es'), use that language
        # - If 'language' key exists but is empty string, keep as empty string (signals auto-detect)
        # - If 'language' key doesn't exist at all, fall back to user's default (backwards compat)
        if 'language' in data:
            # User explicitly chose a language (or auto-detect with empty string)
            language = data.get('language')  # Could be 'es', '', or None
        else:
            # Language not provided - use user's default (backwards compatibility)
            language = recording.owner.transcription_language if recording.owner else None

        min_speakers = data.get('min_speakers') or None
        max_speakers = data.get('max_speakers') or None
        hotwords = (data.get('hotwords') or '').strip() or None
        initial_prompt = (data.get('initial_prompt') or '').strip() or None
        transcription_model = (data.get('transcription_model') or '').strip() or None

        # Convert to int if provided
        if min_speakers:
            try:
                min_speakers = int(min_speakers)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers:
            try:
                max_speakers = int(max_speakers)
            except (ValueError, TypeError):
                max_speakers = None

        # Apply precedence chain mirroring upload_file():
        #   user input > tag defaults > folder defaults > env defaults > user defaults

        # Tag defaults (highest priority after user input). Use the first tag's
        # defaults when present.
        if recording.tags:
            for tag_association in sorted(recording.tag_associations, key=lambda x: x.order):
                tag = tag_association.tag
                if min_speakers is None and tag.default_min_speakers:
                    min_speakers = tag.default_min_speakers
                if max_speakers is None and tag.default_max_speakers:
                    max_speakers = tag.default_max_speakers
                if not hotwords and tag.default_hotwords:
                    hotwords = tag.default_hotwords
                if not initial_prompt and tag.default_initial_prompt:
                    initial_prompt = tag.default_initial_prompt
                if not transcription_model and tag.default_transcription_model:
                    transcription_model = tag.default_transcription_model
                if (min_speakers is not None and max_speakers is not None
                        and hotwords and initial_prompt and transcription_model):
                    break

        # Folder defaults (only if recording has no tags providing the value)
        if recording.folder:
            folder = recording.folder
            if min_speakers is None and folder.default_min_speakers:
                min_speakers = folder.default_min_speakers
            if max_speakers is None and folder.default_max_speakers:
                max_speakers = folder.default_max_speakers
            if not hotwords and folder.default_hotwords:
                hotwords = folder.default_hotwords
            if not initial_prompt and folder.default_initial_prompt:
                initial_prompt = folder.default_initial_prompt
            if not transcription_model and folder.default_transcription_model:
                transcription_model = folder.default_transcription_model

        # Environment variable defaults
        if min_speakers is None and ASR_MIN_SPEAKERS:
            try:
                min_speakers = int(ASR_MIN_SPEAKERS)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers is None and ASR_MAX_SPEAKERS:
            try:
                max_speakers = int(ASR_MAX_SPEAKERS)
            except (ValueError, TypeError):
                max_speakers = None

        # User defaults (lowest priority). The recording owner's account-level
        # hotwords / initial_prompt apply when no tag/folder/env value was set.
        owner = recording.owner
        if owner:
            if not hotwords and owner.transcription_hotwords:
                hotwords = owner.transcription_hotwords
            if not initial_prompt and owner.transcription_initial_prompt:
                initial_prompt = owner.transcription_initial_prompt

        # Validate against admin-curated list and apply admin default when
        # nothing else in the chain set a model. See _resolve_transcription_model.
        transcription_model = _resolve_transcription_model(transcription_model)

        # Enqueue the job with all parameters
        job_params = {
            'language': language,
            'min_speakers': min_speakers,
            'max_speakers': max_speakers,
            'hotwords': hotwords,
            'initial_prompt': initial_prompt,
            'transcription_model': transcription_model,
        }

        job_id = job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='reprocess_transcription',
            params=job_params
        )

        # Get queue position for response
        queue_position = job_queue.get_position_in_queue(recording.id)
        queue_status = job_queue.get_queue_status()

        # Return recording with per-user status and queue info
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Transcription reprocessing queued',
            'recording': recording_dict,
            'queue_position': queue_position,
            'queue_status': queue_status
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reprocessing transcription for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500




@recordings_bp.route('/recording/<int:recording_id>/reprocess_summary', methods=['POST'])
@login_required
def reprocess_summary(recording_id):
    """Reprocess summary for a given recording (requires existing transcription)."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403

        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400

        # Check if transcription is an error message (not actual content)
        if is_transcription_error(recording.transcription):
            return jsonify({'error': 'Cannot generate summary: transcription failed. Please reprocess the transcription first.'}), 400

        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400

        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503

        # Get custom prompt + mode from request if provided
        data = request.get_json() or {}
        custom_prompt = data.get('custom_prompt', '').strip() if data.get('custom_prompt') else None
        prompt_mode = (data.get('prompt_mode') or 'replace').strip().lower()
        if prompt_mode not in ('replace', 'append'):
            prompt_mode = 'replace'
        custom_prompt_append = bool(custom_prompt) and prompt_mode == 'append'

        # Debug logging
        if custom_prompt:
            current_app.logger.info(
                f"Received custom prompt override for recording {recording_id} "
                f"(mode={prompt_mode}, length={len(custom_prompt)})"
            )
        else:
            current_app.logger.info(f"No custom prompt override provided for recording {recording_id}, will use default priority")

        # Per-recording prompt-template variables. Sanitised through the same
        # helper used at upload time so reprocess can't bypass the caps.
        from src.utils.prompt_variables import sanitize_variable_values
        raw_prompt_variables = data.get('prompt_variables')
        if raw_prompt_variables is not None:
            recording.prompt_variables = sanitize_variable_values(raw_prompt_variables)

        # Clear existing summary (status will be set to QUEUED by job_queue.enqueue)
        recording.summary = None

        # Clear existing events since they might be re-extracted during summary generation
        Event.query.filter_by(recording_id=recording_id).delete()

        db.session.commit()

        current_app.logger.info(f"Queueing summary reprocessing for recording {recording_id}" +
                       (f" with custom prompt (length: {len(custom_prompt)})" if custom_prompt else ""))

        # Queue summary generation job
        job_params = {
            'custom_prompt': custom_prompt,
            'custom_prompt_append': custom_prompt_append,
            'user_id': current_user.id
        }
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='reprocess_summary',
            params=job_params
        )

        # Refresh recording to get updated status
        db.session.refresh(recording)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Summary reprocessing started',
            'recording': recording_dict
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reprocessing summary for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/regenerate_title', methods=['POST'])
@login_required
def regenerate_title(recording_id):
    """Regenerate the AI title for a recording based on its existing transcription."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for title generation'}), 400

        if is_transcription_error(recording.transcription):
            return jsonify({'error': 'Cannot generate title: transcription failed. Please reprocess the transcription first.'}), 400

        if client is None:
            return jsonify({'error': 'Title generation service is not available (OpenRouter client not configured)'}), 503

        from src.tasks.processing import _generate_ai_title

        new_title = _generate_ai_title(recording)
        if not new_title:
            return jsonify({'error': 'Failed to generate a title'}), 500

        recording.title = new_title
        db.session.commit()

        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'title': new_title,
            'recording': recording_dict
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error regenerating title for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/recording/<int:recording_id>/reset_status', methods=['POST'])
@login_required
def reset_status(recording_id):
    """Resets the status of a stuck or failed recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        # Allow resetting if it's stuck or failed
        if recording.status in ['PENDING', 'PROCESSING', 'SUMMARIZING', 'FAILED']:
            recording.status = 'FAILED'
            recording.error_message = "Manually reset from stuck or failed state."
            db.session.commit()
            current_app.logger.info(f"Manually reset status for recording {recording_id} to FAILED.")

            # Return recording with per-user status
            recording_dict = recording.to_dict(viewer_user=current_user)
            enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
            return jsonify({'success': True, 'message': 'Recording status has been reset.', 'recording': recording_dict})
        else:
            return jsonify({'error': f'Recording is not in a state that can be reset. Current status: {recording.status}'}), 400

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error resetting status for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

# --- Authentication Routes ---


@recordings_bp.route('/')
@login_required
def index():
    # Check if user is a group admin
    is_team_admin = GroupMembership.query.filter_by(
        user_id=current_user.id,
        role='admin'
    ).first() is not None

    # Pass the ASR config, inquire mode config, and user language preference to the template
    user_language = current_user.ui_language if current_user.is_authenticated and current_user.ui_language else 'en'

    # Calculate if archive toggle should be shown (only when audio-only deletion mode is active)
    enable_archive_toggle = ENABLE_AUTO_DELETION and DELETION_MODE == 'audio_only'

    # Get connector capabilities (new architecture)
    # Defaults to USE_ASR_ENDPOINT for backwards compatibility
    connector_supports_diarization = USE_ASR_ENDPOINT
    connector_supports_speaker_count = USE_ASR_ENDPOINT  # ASR endpoint supports min/max speakers
    connector_supports_hotwords = USE_ASR_ENDPOINT
    connector_supports_initial_prompt = USE_ASR_ENDPOINT
    if USE_NEW_TRANSCRIPTION_ARCHITECTURE:
        try:
            from src.services.transcription import get_registry
            registry = get_registry()
            connector = registry.get_active_connector()
            if connector:
                connector_supports_diarization = connector.supports_diarization
                connector_supports_speaker_count = connector.supports_speaker_count_control
                connector_supports_hotwords = connector.supports_hotwords
                connector_supports_initial_prompt = connector.supports_initial_prompt
        except Exception as e:
            current_app.logger.warning(f"Could not get connector capabilities: {e}")

    # Phase B of #287 (c)(d): server-side chunk streaming for in-app
    # recordings. Off by default; enabling the env var flips the
    # data-server-recording-chunks attribute and the audio composable
    # opens an /upload/session for each in-app recording instead of
    # accumulating the whole blob in RAM.
    server_recording_chunks_enabled = os.environ.get(
        'ENABLE_SERVER_RECORDING_CHUNKS', 'false'
    ).lower() == 'true'

    # Absolute hours ceiling on in-app recordings (Phase C of #287 c/d).
    # Replaces the legacy hard-stop-on-size behaviour for users with
    # server-side streaming enabled; size-based stop still applies as a
    # fallback when streaming is off.
    try:
        recording_max_hours = float(os.environ.get('RECORDING_MAX_HOURS', '8'))
    except (TypeError, ValueError):
        recording_max_hours = 8.0

    # MediaRecorder timeslice (seconds) for in-app recordings: how often a
    # chunk is emitted and streamed to the server. Smaller = finer crash
    # recovery but more requests / files / DB commits; larger = less server
    # load. Default 5s. At scale (many concurrent recorders on SQLite) raise
    # this and/or RECORDING_SESSION_COMMIT_BATCH_SIZE. Clamped to [1, 60].
    try:
        recording_chunk_seconds = int(float(os.environ.get('RECORDING_CHUNK_SECONDS', '5')))
    except (TypeError, ValueError):
        recording_chunk_seconds = 5
    recording_chunk_seconds = max(1, min(60, recording_chunk_seconds))

    return render_template('index.html',
                         use_asr_endpoint=USE_ASR_ENDPOINT,  # Backwards compat
                         connector_supports_diarization=connector_supports_diarization,
                         connector_supports_speaker_count=connector_supports_speaker_count,
                         connector_supports_hotwords=connector_supports_hotwords,
                         connector_supports_initial_prompt=connector_supports_initial_prompt,
                         inquire_mode_enabled=ENABLE_INQUIRE_MODE,
                         enable_archive_toggle=enable_archive_toggle,
                         enable_internal_sharing=ENABLE_INTERNAL_SHARING,
                         user_language=user_language,
                         is_team_admin=is_team_admin,
                         server_recording_chunks_enabled=server_recording_chunks_enabled,
                         recording_max_hours=recording_max_hours,
                         recording_chunk_seconds=recording_chunk_seconds)



def get_accessible_recording_ids(user_id):
    """
    Get all recording IDs that a user has access to.

    Includes:
    - Recordings owned by the user
    - Recordings shared with the user via InternalShare
    - Recordings shared via group tags (if team membership exists)

    Args:
        user_id (int): User ID to check access for

    Returns:
        list: List of recording IDs the user can access
    """
    accessible_ids = set()

    # 1. User's own recordings
    own_recordings = db.session.query(Recording.id).filter_by(user_id=user_id).all()
    accessible_ids.update([r.id for r in own_recordings])

    # 2. Internally shared recordings
    if ENABLE_INTERNAL_SHARING:
        shared_recordings = db.session.query(InternalShare.recording_id).filter_by(
            shared_with_user_id=user_id
        ).all()
        accessible_ids.update([r.recording_id for r in shared_recordings])

    return list(accessible_ids)


@recordings_bp.route('/recordings', methods=['GET'])
def get_recordings():
    """Get all recordings for the current user (simple list)."""
    try:
        # Check if user is logged in
        if not current_user.is_authenticated:
            return jsonify([])  # Return empty array if not logged in

        # Filter recordings by the current user
        stmt = select(Recording).where(Recording.user_id == current_user.id).order_by(Recording.created_at.desc())
        recordings = db.session.execute(stmt).scalars().all()
        # Pre-batch the duplicate-info groupings so we don't run a
        # per-row query inside each recording's to_dict() call.
        dup_map = Recording.get_duplicate_info_map(current_user.id)
        return jsonify([
            recording.to_dict(viewer_user=current_user, duplicate_info_map=dup_map)
            for recording in recordings
        ])
    except Exception as e:
        current_app.logger.error(f"Error fetching recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recordings', methods=['GET'])
@login_required
def get_recordings_paginated():
    """Get recordings with pagination and server-side filtering (includes shared recordings)."""
    import re
    try:
        # Parse query parameters
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 25, type=int), 100)  # Cap at 100 per page
        search_query = request.args.get('q', '').strip()
        show_archived = request.args.get('archived', '').lower() == 'true'
        show_shared = request.args.get('shared', '').lower() == 'true'
        show_starred = request.args.get('starred', '').lower() == 'true'
        show_inbox = request.args.get('inbox', '').lower() == 'true'
        sort_by = request.args.get('sort_by', 'created_at')  # 'created_at' or 'meeting_date'
        folder_filter = request.args.get('folder', '').strip()  # folder_id or 'none' for no folder

        # Get all accessible recording IDs (own + shared)
        accessible_recording_ids = get_accessible_recording_ids(current_user.id)

        if not accessible_recording_ids:
            return jsonify({
                'recordings': [],
                'pagination': {
                    'page': page,
                    'per_page': per_page,
                    'total': 0,
                    'total_pages': 0,
                    'has_next': False,
                    'has_prev': False
                }
            })

        # Build base query to include accessible recordings
        stmt = select(Recording).where(Recording.id.in_(accessible_recording_ids))

        # Apply archived filter (AND with other filters)
        if show_archived:
            # Only show recordings where audio has been deleted
            stmt = stmt.where(Recording.audio_deleted_at.is_not(None))

        # Apply shared filter (AND with other filters)
        if show_shared:
            # Only show recordings shared with current user (not owned by them)
            stmt = stmt.where(Recording.user_id != current_user.id)

        # Apply starred filter (AND with other filters)
        # For starred/inbox we need to consider both owned recordings and shared recordings
        if show_starred:
            from src.models.sharing import SharedRecordingState
            # For owned recordings: check Recording.is_highlighted
            # For shared recordings: check SharedRecordingState.is_highlighted
            starred_subq = select(SharedRecordingState.recording_id).where(
                db.and_(
                    SharedRecordingState.user_id == current_user.id,
                    SharedRecordingState.is_highlighted == True
                )
            ).scalar_subquery()
            stmt = stmt.where(
                db.or_(
                    db.and_(Recording.user_id == current_user.id, Recording.is_highlighted == True),
                    Recording.id.in_(starred_subq)
                )
            )

        # Apply inbox filter (AND with other filters)
        if show_inbox:
            from src.models.sharing import SharedRecordingState
            # For owned recordings: check Recording.is_inbox
            # For shared recordings: check SharedRecordingState.is_inbox
            inbox_subq = select(SharedRecordingState.recording_id).where(
                db.and_(
                    SharedRecordingState.user_id == current_user.id,
                    SharedRecordingState.is_inbox == True
                )
            ).scalar_subquery()
            stmt = stmt.where(
                db.or_(
                    db.and_(Recording.user_id == current_user.id, Recording.is_inbox == True),
                    Recording.id.in_(inbox_subq)
                )
            )

        # Apply folder filter (AND with other filters)
        if folder_filter:
            if folder_filter.lower() == 'none':
                # Filter recordings with no folder
                stmt = stmt.where(Recording.folder_id.is_(None))
            else:
                # Filter by specific folder_id
                try:
                    folder_id = int(folder_filter)
                    stmt = stmt.where(Recording.folder_id == folder_id)
                except ValueError:
                    pass  # Invalid folder_id, ignore filter

        # Apply search filters if provided
        if search_query:
            # Extract date filters
            date_filters = re.findall(r'date:(\S+)', search_query.lower())
            date_from_filters = re.findall(r'date_from:(\S+)', search_query.lower())
            date_to_filters = re.findall(r'date_to:(\S+)', search_query.lower())
            tag_filters = re.findall(r'tag:(\S+)', search_query.lower())
            speaker_filters = re.findall(r'speaker:(\S+)', search_query.lower())

            # Remove special syntax to get text search
            text_query = re.sub(r'date:\S+', '', search_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_from:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_to:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'tag:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'speaker:\S+', '', text_query, flags=re.IGNORECASE).strip()

            # Apply date filters
            for date_filter in date_filters:
                if date_filter == 'today':
                    today = datetime.now().date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == today,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == today
                            )
                        )
                    )
                elif date_filter == 'yesterday':
                    yesterday = datetime.now().date() - timedelta(days=1)
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == yesterday,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == yesterday
                            )
                        )
                    )
                elif date_filter == 'thisweek':
                    today = datetime.now().date()
                    start_of_week = today - timedelta(days=today.weekday())
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_week,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_week
                            )
                        )
                    )
                elif date_filter == 'lastweek':
                    today = datetime.now().date()
                    end_of_last_week = today - timedelta(days=today.weekday())
                    start_of_last_week = end_of_last_week - timedelta(days=7)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= start_of_last_week,
                                Recording.meeting_date < end_of_last_week
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_last_week,
                                db.func.date(Recording.created_at) < end_of_last_week
                            )
                        )
                    )
                elif date_filter == 'thismonth':
                    today = datetime.now().date()
                    start_of_month = today.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_month,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_month
                            )
                        )
                    )
                elif date_filter == 'lastmonth':
                    today = datetime.now().date()
                    first_day_this_month = today.replace(day=1)
                    last_day_last_month = first_day_this_month - timedelta(days=1)
                    first_day_last_month = last_day_last_month.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= first_day_last_month,
                                Recording.meeting_date <= last_day_last_month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= first_day_last_month,
                                db.func.date(Recording.created_at) <= last_day_last_month
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', date_filter):
                    # Specific date format YYYY-MM-DD
                    target_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == target_date,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == target_date
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}$', date_filter):
                    # Month format YYYY-MM
                    year, month = map(int, date_filter.split('-'))
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                db.extract('year', Recording.meeting_date) == year,
                                db.extract('month', Recording.meeting_date) == month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year,
                                db.extract('month', Recording.created_at) == month
                            )
                        )
                    )
                elif re.match(r'^\d{4}$', date_filter):
                    # Year format YYYY
                    year = int(date_filter)
                    stmt = stmt.where(
                        db.or_(
                            db.extract('year', Recording.meeting_date) == year,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year
                            )
                        )
                    )

            # Apply date range filters
            if date_from_filters and date_from_filters[0]:
                try:
                    date_from = datetime.strptime(date_from_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= date_from,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= date_from
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore

            if date_to_filters and date_to_filters[0]:
                try:
                    date_to = datetime.strptime(date_to_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date <= date_to,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) <= date_to
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore

            # Apply tag filters
            if tag_filters:
                # Join with tags table and filter by tag names
                tag_conditions = []
                for tag_filter in tag_filters:
                    # Replace underscores back to spaces for matching
                    tag_name = tag_filter.replace('_', ' ')
                    tag_conditions.append(Tag.name.ilike(f'%{tag_name}%'))

                stmt = stmt.join(RecordingTag).join(Tag).where(db.or_(*tag_conditions))

            # Apply speaker filters
            if speaker_filters:
                speaker_conditions = []
                for speaker_filter in speaker_filters:
                    # Replace underscores back to spaces for matching
                    speaker_name = speaker_filter.replace('_', ' ')
                    speaker_conditions.append(Recording.participants.ilike(f'%{speaker_name}%'))
                stmt = stmt.where(db.or_(*speaker_conditions))

            # Apply text search
            if text_query:
                from src.models.sharing import SharedRecordingState

                # Search in user-specific notes:
                # - For owned recordings: search Recording.notes
                # - For shared recordings: search SharedRecordingState.personal_notes

                text_conditions = [
                    Recording.title.ilike(f'%{text_query}%'),
                    Recording.participants.ilike(f'%{text_query}%'),
                    Recording.transcription.ilike(f'%{text_query}%'),
                    # Search owner's notes for owned recordings
                    db.and_(
                        Recording.user_id == current_user.id,
                        Recording.notes.ilike(f'%{text_query}%')
                    )
                ]

                # Add search for personal notes in shared recordings
                # Use a subquery to check if personal_notes match
                shared_notes_subq = select(SharedRecordingState.recording_id).where(
                    db.and_(
                        SharedRecordingState.user_id == current_user.id,
                        SharedRecordingState.personal_notes.ilike(f'%{text_query}%')
                    )
                ).scalar_subquery()

                text_conditions.append(Recording.id.in_(shared_notes_subq))

                stmt = stmt.where(db.or_(*text_conditions))

        # Apply ordering based on sort_by parameter
        if sort_by == 'meeting_date':
            # Sort by meeting_date first, fall back to created_at if no meeting_date
            stmt = stmt.order_by(
                db.case(
                    (Recording.meeting_date.is_not(None), Recording.meeting_date),
                    else_=db.func.date(Recording.created_at)
                ).desc(),
                Recording.created_at.desc()
            )
        else:
            # Default: sort by created_at (upload/processing date)
            stmt = stmt.order_by(Recording.created_at.desc())

        # Get total count for pagination info
        count_stmt = select(db.func.count()).select_from(stmt.subquery())
        total_count = db.session.execute(count_stmt).scalar()

        # Apply pagination
        offset = (page - 1) * per_page
        stmt = stmt.offset(offset).limit(per_page)

        # Execute query
        recordings = db.session.execute(stmt).scalars().all()

        # Enrich recordings with sharing metadata
        enriched_recordings = []
        for recording in recordings:
            rec_dict = recording.to_list_dict(viewer_user=current_user)

            # Add sharing metadata
            is_owner = recording.user_id == current_user.id
            rec_dict['is_owner'] = is_owner

            # Get per-user status (owner uses Recording fields, recipients use SharedRecordingState)
            user_inbox, user_highlighted = get_user_recording_status(recording, current_user)
            rec_dict['is_inbox'] = user_inbox
            rec_dict['is_highlighted'] = user_highlighted

            # Add edit permission info (uses has_recording_access which checks group admin status)
            rec_dict['can_edit'] = has_recording_access(recording, current_user, require_edit=True)

            # Add delete permission info (only owner can delete)
            rec_dict['can_delete'] = is_owner and (USERS_CAN_DELETE or current_user.is_admin)

            if not is_owner:
                # This is a shared recording - get owner info and share permissions
                owner = db.session.get(User, recording.user_id)
                rec_dict['owner_username'] = owner.username if owner else "Unknown"
                rec_dict['is_shared'] = True
                # Don't show outgoing share count for recordings you don't own
                rec_dict['shared_with_count'] = 0
                rec_dict['public_share_count'] = 0

                # Get share permissions
                share = InternalShare.query.filter_by(
                    recording_id=recording.id,
                    shared_with_user_id=current_user.id
                ).first()

                if share:
                    rec_dict['share_info'] = {
                        'share_id': share.id,
                        'owner_username': owner.username if owner else "Unknown",
                        'can_edit': share.can_edit,
                        'can_reshare': share.can_reshare,
                        'shared_at': share.created_at.isoformat()
                    }
                else:
                    # Fallback if share record not found (shouldn't happen)
                    rec_dict['share_info'] = {
                        'can_edit': False,
                        'can_reshare': False
                    }
            else:
                rec_dict['is_shared'] = False

            # Check if recording has group tags (among visible tags)
            visible_tags = recording.get_visible_tags(current_user)
            has_group_tags = any(tag.is_group_tag for tag in visible_tags)
            rec_dict['has_group_tags'] = has_group_tags

            enriched_recordings.append(rec_dict)

        # Calculate pagination metadata
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_prev = page > 1

        return jsonify({
            'recordings': enriched_recordings,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_prev': has_prev
            }
        })

    except Exception as e:
        current_app.logger.error(f"Error fetching paginated recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/save', methods=['POST'])
@login_required
def save_metadata():
    """Save recording metadata (title, participants, notes, summary, etc.)."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        recording_id = data.get('id')
        if not recording_id:
            return jsonify({'error': 'No recording ID provided'}), 400

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has at least view access
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Handle notes separately - no edit permission required (user-specific)
        if 'notes' in data:
            if recording.user_id == current_user.id:
                # Owner saves to Recording.notes
                recording.notes = sanitize_html(data['notes']) if data['notes'] else data['notes']
            else:
                # Shared user saves to personal_notes (requires SharedRecordingState)
                from src.models.sharing import SharedRecordingState
                state = SharedRecordingState.query.filter_by(
                    recording_id=recording.id,
                    user_id=current_user.id
                ).first()

                if not state:
                    # Create SharedRecordingState if it doesn't exist
                    state = SharedRecordingState(
                        recording_id=recording.id,
                        user_id=current_user.id,
                        is_inbox=True,
                        is_highlighted=False
                    )
                    db.session.add(state)

                state.personal_notes = sanitize_html(data['notes']) if data['notes'] else data['notes']

        # Determine if any fields requiring edit permission are being updated
        edit_fields = ['title', 'participants', 'summary', 'meeting_date']
        requires_edit = any(field in data for field in edit_fields)

        # If edit fields are present, check for edit permission
        if requires_edit and not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        # Update fields requiring edit permission
        if requires_edit:
            if 'title' in data:
                recording.title = data['title']
            if 'participants' in data:
                recording.participants = data['participants']
            if 'summary' in data:
                recording.summary = sanitize_html(data['summary']) if data['summary'] else data['summary']
            if 'meeting_date' in data:
                try:
                    date_str = data['meeting_date']
                    if date_str:
                        # Try to parse as full ISO datetime first
                        try:
                            recording.meeting_date = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                        except (ValueError, AttributeError):
                            # Fall back to date-only format, preserve existing time if available
                            parsed_date = datetime.strptime(date_str, '%Y-%m-%d')
                            if recording.meeting_date:
                                # Preserve existing time
                                existing_time = recording.meeting_date.time()
                                recording.meeting_date = datetime.combine(parsed_date.date(), existing_time)
                            else:
                                # No existing time, use the parsed date with midnight time
                                recording.meeting_date = parsed_date
                    else:
                        recording.meeting_date = None
                except (ValueError, TypeError) as e:
                    current_app.logger.warning(f"Could not parse meeting_date '{data.get('meeting_date')}': {e}")

        # Handle per-user status fields (only requires view permission)
        if 'is_inbox' in data or 'is_highlighted' in data:
            set_user_recording_status(
                recording,
                current_user,
                is_inbox=data.get('is_inbox'),
                is_highlighted=data.get('is_highlighted')
            )

        db.session.commit()

        # Re-export the recording if auto-export is enabled and editable fields were changed
        if requires_edit:
            export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({'success': True, 'recording': recording_dict})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving metadata for recording {data.get('id')}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while saving.'}), 500


@recordings_bp.route('/recording/<int:recording_id>/update_transcription', methods=['POST'])
@login_required
def update_transcription(recording_id):
    """Updates the transcription content for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        new_transcription = data.get('transcription')

        if new_transcription is None:
            return jsonify({'error': 'No transcription data provided'}), 400

        # The incoming data could be a JSON string (from ASR edit) or plain text
        recording.transcription = new_transcription

        # Optional: If the transcription changes, we might want to indicate that the summary is outdated.
        # For now, we'll just save the transcript. A "regenerate summary" button could be a good follow-up.

        db.session.commit()
        current_app.logger.info(f"Transcription for recording {recording_id} was updated.")

        # Re-export the recording if auto-export is enabled
        export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({'success': True, 'message': 'Transcription updated successfully.', 'recording': recording_dict})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating transcription for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while updating the transcription.'}), 500

# Toggle inbox status endpoint


@recordings_bp.route('/recording/<int:recording_id>/toggle_inbox', methods=['POST'])
@login_required
def toggle_inbox(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Only require view access (not edit) - users can manage their own inbox status
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Get current status and toggle it
        current_inbox, current_highlighted = get_user_recording_status(recording, current_user)
        new_inbox, new_highlighted = set_user_recording_status(recording, current_user, is_inbox=not current_inbox)

        return jsonify({'success': True, 'is_inbox': new_inbox})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling inbox status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Toggle highlighted status endpoint


@recordings_bp.route('/recording/<int:recording_id>/toggle_highlight', methods=['POST'])
@login_required
def toggle_highlight(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Only require view access (not edit) - users can manage their own highlight status
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Get current status and toggle it
        current_inbox, current_highlighted = get_user_recording_status(recording, current_user)
        new_inbox, new_highlighted = set_user_recording_status(recording, current_user, is_highlighted=not current_highlighted)

        return jsonify({'success': True, 'is_highlighted': new_highlighted})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling highlighted status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500




@recordings_bp.route('/share-target', methods=['POST', 'GET'])
@login_required
def share_target():
    """PWA Web Share Target receiver (issue #285).

    Android Chrome / iOS Safari 16.4+ POST a multipart/form-data body to
    this endpoint when the user picks PXE MeetingMitra from the native share sheet.
    The file arrives under the field name declared in the manifest
    (``shared_audio``). Any text/title/url params the share sheet sent
    become the recording's notes.

    CSRF is exempted at app startup (the share sheet has no way to
    round-trip a CSRF token). Authentication still happens via the
    session cookie carried by the browser; unauthenticated visitors hit
    @login_required and are bounced to the login page.

    Implementation: save the file to ``UPLOAD_FOLDER``, create a Recording
    row applying the same user / folder / tag defaults as ``upload_file``
    where they are unambiguous (no tags / folder context in a share),
    then enqueue a transcription job through the existing job queue.
    Redirect the browser to the SPA so the user lands on the recordings
    list with the new row visible.
    """
    if request.method == 'GET':
        # Some platforms pre-flight with GET. Just land the user on the SPA.
        return redirect(url_for('recordings.index'))

    shared = request.files.get('shared_audio') or request.files.get('file')
    if not shared or not getattr(shared, 'filename', ''):
        return redirect(url_for('recordings.index') + '?share_target_error=missing_file')

    original_filename = shared.filename
    safe_filename = secure_filename(original_filename) or 'shared-recording.webm'
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{timestamp}_{safe_filename}")

    try:
        shared.save(filepath)
    except RequestEntityTooLarge:
        return redirect(url_for('recordings.index') + '?share_target_error=too_large')
    except Exception as save_err:
        current_app.logger.warning(f"share-target file save failed: {save_err}")
        return redirect(url_for('recordings.index') + '?share_target_error=save_failed')

    file_size = os.path.getsize(filepath)

    # Build the recording row. Tags and folder cannot be inferred from a
    # share sheet, so we leave them empty; the user can adjust from the
    # recording detail view. Title comes from the share sheet's title field
    # if present; otherwise we use the SAME placeholder a normal upload gets
    # (resolve_upload_title) so the AI title task recognises it and generates
    # a title — previously the filename stem was used, which the title task
    # treated as a user-chosen title and skipped, leaving shared files untitled.
    from src.utils.titles import resolve_upload_title
    share_title = resolve_upload_title(request.form.get('title'), original_filename)

    notes_parts = []
    for key in ('title', 'text', 'url'):
        val = (request.form.get(key) or '').strip()
        if val:
            notes_parts.append(val)
    share_notes = '\n\n'.join(notes_parts) if notes_parts else None

    recording = Recording(
        audio_path=filepath,
        original_filename=original_filename,
        title=share_title,
        status='PENDING',
        user_id=current_user.id,
        notes=share_notes,
        file_size=file_size,
    )
    db.session.add(recording)
    db.session.commit()

    # Enqueue transcription with user's defaults. Mirrors the precedence
    # chain in upload_file for the fields that apply when no tag/folder
    # context is available.
    job_params = {
        'language': current_user.transcription_language,
        'hotwords': current_user.transcription_hotwords,
        'initial_prompt': current_user.transcription_initial_prompt,
    }
    try:
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='transcribe',
            params=job_params,
            is_new_upload=True,
        )
    except Exception as queue_err:
        current_app.logger.warning(f"share-target enqueue failed for recording {recording.id}: {queue_err}")
        recording.status = 'FAILED'
        recording.transcription = f"Processing failed: {queue_err}"
        db.session.commit()
        return redirect(url_for('recordings.index') + '?share_target_error=queue_failed')

    return redirect(url_for('recordings.index') + f'?share_target=ok&recording_id={recording.id}')


@recordings_bp.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        storage = get_storage_service()
        staging_dir = storage.get_staging_dir()
        filepath = os.path.join(staging_dir, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_filename}")

        # Per-upload override for VIDEO_RETENTION. When True, the
        # processing pipeline extracts audio and discards the video stream
        # regardless of the server's VIDEO_RETENTION env var. Also signals
        # that this upload is allowed to exceed max_file_size_mb because
        # only the extracted audio needs to fit that limit. Sent by the
        # frontend either explicitly (toggle) or implicitly when
        # VIDEO_RETENTION is off and the queued file is video.
        keep_audio_only_flag = request.form.get('keep_audio_only', 'false').lower() == 'true'
        effective_audio_only = keep_audio_only_flag or not VIDEO_RETENTION

        # Detect "this looks like a video" by extension up front so the
        # size-gate can pick the right limit before reading any bytes.
        # ffprobe runs later and is authoritative for the keep-video
        # decision; the extension list here is a pre-save heuristic.
        _VIDEO_EXTS = {'.mp4', '.mov', '.mkv', '.avi', '.webm', '.m4v', '.wmv', '.flv', '.ts', '.mts'}
        _ext_lower = os.path.splitext(original_filename)[1].lower()
        is_likely_video_by_ext = _ext_lower in _VIDEO_EXTS

        # Get original file size
        file.seek(0, os.SEEK_END)
        original_file_size = file.tell()
        file.seek(0)

        # Resolve effective size limit. Video files (by extension) get
        # the larger video cap regardless of whether the user opted
        # into "Keep audio only". Audio files use the regular cap.
        # This is the option-B model: admins set a single per-file-type
        # upload cap and don't have to think about how the keep-audio
        # toggle interacts with size validation. The post-extraction
        # guard further down still enforces that the *stored* artifact
        # fits the regular limit when only audio is kept.
        regular_limit_mb = int(SystemSetting.get_setting('max_file_size_mb', 250))
        audio_only_limit_mb = int(
            SystemSetting.get_setting('max_audio_only_video_size_mb', regular_limit_mb * 4)
        )
        if is_likely_video_by_ext:
            effective_limit_mb = audio_only_limit_mb
        else:
            effective_limit_mb = regular_limit_mb
        effective_limit_bytes = effective_limit_mb * 1024 * 1024

        # Check size limit before saving - only enforce if chunking is disabled or connector handles it
        max_content_length = current_app.config.get('MAX_CONTENT_LENGTH')

        # Get connector specifications for chunking decisions
        connector_specs = None
        if USE_NEW_TRANSCRIPTION_ARCHITECTURE:
            try:
                from src.services.transcription import get_registry
                registry = get_registry()
                connector = registry.get_active_connector()
                if connector:
                    connector_specs = connector.specifications
            except Exception as e:
                current_app.logger.warning(f"Could not get connector specs for upload: {e}")

        # Skip size check if chunking is enabled (app-level or connector handles internally)
        should_enforce_size_limit = True
        if chunking_service:
            from src.audio_chunking import get_effective_chunking_config
            chunking_config = get_effective_chunking_config(connector_specs)
            if chunking_config.enabled or chunking_config.source == 'connector_internal':
                should_enforce_size_limit = False
                if chunking_config.source == 'connector_internal':
                    current_app.logger.info(f"Connector handles chunking internally - skipping {original_file_size/1024/1024:.1f}MB size limit check")
                elif chunking_config.mode == 'size':
                    current_app.logger.info(f"Size-based chunking enabled ({chunking_config.limit_value}MB, source={chunking_config.source}) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
                else:
                    current_app.logger.info(f"Duration-based chunking enabled ({chunking_config.limit_value}s, source={chunking_config.source}) - skipping {original_file_size/1024/1024:.1f}MB size limit check")

        # Two ceilings to consider. The Werkzeug `max_content_length` is
        # the WSGI hard cap; the smaller `effective_limit_bytes` is the
        # per-request policy cap (which differs between regular and
        # audio-only-video uploads). Reject if EITHER fires.
        if should_enforce_size_limit and (
            (max_content_length and original_file_size > max_content_length)
            or original_file_size > effective_limit_bytes
        ):
            raise RequestEntityTooLarge()

        file.save(filepath)
        current_app.logger.info(f"File saved to {filepath}")

        # Compute file hash on the ORIGINAL upload before any conversion/compression.
        # Lossy re-encoding (e.g. FLAC→MP3) produces different bytes each run,
        # so hashing after conversion would miss duplicates.
        file_hash = None
        duplicate_warning = None
        try:
            file_hash = compute_file_sha256(filepath)
            existing = Recording.query.filter_by(
                user_id=current_user.id, file_hash=file_hash
            ).first()
            if existing:
                duplicate_warning = {
                    'existing_recording_id': existing.id,
                    'existing_title': existing.title,
                    'existing_created_at': existing.created_at.isoformat() if existing.created_at else None
                }
                current_app.logger.info(
                    f"Duplicate file detected for user {current_user.id}: "
                    f"hash={file_hash[:12]}... matches recording {existing.id}"
                )
        except Exception as e:
            current_app.logger.warning(f"Could not compute file hash: {e}")

        # --- Convert files only when chunking is needed ---
        filename_lower = original_filename.lower()

        # Check if chunking will be needed for this file (uses connector-aware logic)
        needs_chunking_for_processing = bool(
            chunking_service and
            chunking_service.needs_chunking(filepath, USE_ASR_ENDPOINT, connector_specs)
        )

        # Probe once and use shared conversion utility
        # Scale timeout based on file size — large files (especially MP4 with moov at end) need more time
        file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        probe_timeout = max(10, min(60, int(file_size_mb / 10)))  # 10s min, scales ~1s per 10MB, 60s max
        codec_info = None
        try:
            codec_info = get_codec_info(filepath, timeout=probe_timeout)
            current_app.logger.info(
                f"Detected codec for {original_filename}: "
                f"audio_codec={codec_info.get('audio_codec')}, "
                f"has_video={codec_info.get('has_video', False)}"
            )
        except FFProbeError as e:
            current_app.logger.warning(f"Failed to probe {original_filename} (timeout={probe_timeout}s): {e}. Will attempt conversion.")
            codec_info = None

        # Video retention/passthrough: skip conversion for videos, processing pipeline handles extraction
        has_video = codec_info.get('has_video', False) if codec_info else False

        # Fallback: if probe failed but VIDEO_RETENTION or VIDEO_PASSTHROUGH_ASR is on, check file extension
        # to avoid silently discarding video from files we couldn't probe
        if codec_info is None and (VIDEO_RETENTION or VIDEO_PASSTHROUGH_ASR) and not has_video:
            video_extensions = {'.mp4', '.mov', '.mkv', '.avi', '.webm', '.m4v', '.wmv', '.flv', '.ts', '.mts'}
            file_ext = os.path.splitext(original_filename)[1].lower()
            if file_ext in video_extensions:
                has_video = True
                current_app.logger.info(
                    f"Probe failed but file extension '{file_ext}' indicates video — "
                    f"treating as video for {'VIDEO_PASSTHROUGH_ASR' if VIDEO_PASSTHROUGH_ASR else 'VIDEO_RETENTION'}"
                )
        # Per-upload `keep_audio_only` overrides VIDEO_RETENTION for this
        # request. VIDEO_PASSTHROUGH_ASR is an admin escape hatch that
        # sends video directly to the ASR endpoint — left untouched here.
        keep_video_for_this_upload = (
            ((VIDEO_RETENTION and not keep_audio_only_flag) or VIDEO_PASSTHROUGH_ASR)
            and has_video
        )
        if keep_video_for_this_upload:
            current_app.logger.info(f"Video {'passthrough' if VIDEO_PASSTHROUGH_ASR else 'retention'}: keeping original video, skipping conversion")
        else:
            # Use shared conversion utility - handles ALL conversion needs (codec conversion + compression)
            try:
                result = convert_if_needed(
                    filepath,
                    original_filename=original_filename,
                    codec_info=codec_info,
                    needs_chunking=needs_chunking_for_processing,
                    is_asr_endpoint=USE_ASR_ENDPOINT,
                    delete_original=True,
                    connector_specs=connector_specs  # Pass connector specs for codec restrictions
                )
                filepath = result.output_path

                # Log what happened
                if result.was_converted:
                    current_app.logger.info(f"File converted: {result.original_codec} -> {result.final_codec}")
                if result.was_compressed:
                    current_app.logger.info(f"File compressed: {result.size_reduction_percent:.1f}% size reduction")

            except FFmpegNotFoundError as e:
                current_app.logger.error(f"FFmpeg not found: {e}")
                return jsonify({'error': 'Audio conversion tool (FFmpeg) not found on server.'}), 500
            except FFmpegError as e:
                current_app.logger.error(f"FFmpeg conversion failed for {filepath}: {e}")
                return jsonify({'error': f'Failed to convert audio file: {str(e)}'}), 500

        # Get final file size (of original or converted file)
        final_file_size = os.path.getsize(filepath)

        # Post-extraction size guard. When the upload exceeded the regular
        # max_file_size_mb under the "audio-only video" exception, the
        # stored file (after audio extraction) must still fit the regular
        # limit. If it does not, the user picked a video whose audio track
        # alone is enormous AND chunking is off; reject the upload,
        # clean up, and let the user know. When chunking is on the
        # large extracted audio is fine — the chunking pipeline will
        # split it for the ASR call.
        chunking_will_handle_large_audio = (
            chunking_service is not None and not should_enforce_size_limit
        )
        if (
            effective_audio_only
            and is_likely_video_by_ext
            and final_file_size > regular_limit_mb * 1024 * 1024
            and not chunking_will_handle_large_audio
        ):
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
            except OSError:
                pass
            current_app.logger.warning(
                f"Audio-only extraction left {final_file_size/1024/1024:.1f}MB of audio, "
                f"which exceeds regular limit {regular_limit_mb}MB and chunking is off. "
                f"Rejecting upload."
            )
            return jsonify({
                'error': (
                    f'The extracted audio is {final_file_size / 1024 / 1024:.0f} MB, '
                    f'larger than the {regular_limit_mb} MB stored-file limit. '
                    f'Enable chunking in admin settings (ENABLE_CHUNKING=true) or '
                    f'use a shorter / lower-bitrate source video.'
                ),
                'max_size_mb': float(regular_limit_mb),
                'extracted_audio_mb': round(final_file_size / 1024 / 1024, 1),
                'audio_only_mode': True,
                'chunking_enabled': False,
            }), 413

        # (file_hash and duplicate_warning already computed above, before conversion)

        # Determine MIME type of the final file
        mime_type, _ = mimetypes.guess_type(filepath)

        # For a retained-video file, derive the MIME from the file's ACTUAL
        # container via the shared probe-driven resolver (the same one the
        # processing task uses) rather than the extension guess, which
        # mislabels ambiguous containers — notably '.webm', registered as
        # audio/webm in app.py — and would hide the video player in the UI.
        # We already have codec_info from the probe above, so no second probe.
        if keep_video_for_this_upload:
            from src.utils.mime import resolve_media_mime
            corrected = resolve_media_mime(filepath, codec_info=codec_info, has_video=True)
            if corrected != mime_type:
                current_app.logger.info(f"Resolved video mime {mime_type!r} -> {corrected!r} from container for {filepath}")
            mime_type = corrected

        current_app.logger.info(f"Final MIME type: {mime_type} for file {filepath}")

        # Cache media duration while the upload is still available locally.
        # This avoids later S3 materialization just to serialize API responses.
        audio_duration_seconds = None
        try:
            detected_duration = get_duration(filepath, timeout=30)
            if detected_duration is not None and detected_duration > 0:
                audio_duration_seconds = float(detected_duration)
        except Exception as e:
            current_app.logger.warning(f"Could not determine duration for upload {original_filename}: {e}")

        # Get notes from the form
        notes = request.form.get('notes')

        # Get optional user-provided title and meeting_date
        user_title = request.form.get('title')
        user_meeting_date = request.form.get('meeting_date')

        # Get file's lastModified timestamp from client (milliseconds since epoch)
        file_last_modified = request.form.get('file_last_modified')

        # Get selected tags if provided (multiple tags support)
        selected_tags = []
        tag_index = 0
        while True:
            tag_id_key = f'tag_ids[{tag_index}]'
            tag_id = request.form.get(tag_id_key)
            if not tag_id:
                break

            # Check if tag belongs to user OR is a group tag where user is a member
            tag = Tag.query.filter_by(id=tag_id).first()
            if tag:
                # Allow tag if it's user's own tag OR it's a group tag where user is a member
                if tag.user_id == current_user.id or (tag.group_id and GroupMembership.query.filter_by(group_id=tag.group_id, user_id=current_user.id).first()):
                    selected_tags.append(tag)
            tag_index += 1

        # For backward compatibility with single tag uploads
        if not selected_tags:
            single_tag_id = request.form.get('tag_id')
            if single_tag_id:
                # Check if tag belongs to user OR is a group tag where user is a member
                tag = Tag.query.filter_by(id=single_tag_id).first()
                if tag and (tag.user_id == current_user.id or (tag.group_id and GroupMembership.query.filter_by(group_id=tag.group_id, user_id=current_user.id).first())):
                    selected_tags.append(tag)

        # Get folder_id if provided
        selected_folder = None
        folder_id = request.form.get('folder_id')
        if folder_id:
            folder = Folder.query.filter_by(id=folder_id).first()
            if folder:
                # Allow folder if it's user's own folder OR it's a group folder where user is a member
                if folder.user_id == current_user.id or (folder.group_id and GroupMembership.query.filter_by(group_id=folder.group_id, user_id=current_user.id).first()):
                    selected_folder = folder

        # Get ASR advanced options if provided
        language = request.form.get('language', '')
        min_speakers = request.form.get('min_speakers') or None
        max_speakers = request.form.get('max_speakers') or None
        hotwords = request.form.get('hotwords', '').strip() or None
        initial_prompt = request.form.get('initial_prompt', '').strip() or None
        transcription_model = request.form.get('transcription_model', '').strip() or None

        # Per-recording prompt-template variables. Sent as a JSON string from
        # the upload form so multiple values fit in a single form field. The
        # sanitiser enforces identifier-shaped keys and per-value / total
        # size caps so an arbitrary JSON blob cannot bloat the column.
        from src.utils.prompt_variables import sanitize_variable_values
        prompt_variables = None
        raw_prompt_variables = request.form.get('prompt_variables')
        if raw_prompt_variables:
            try:
                parsed_vars = json.loads(raw_prompt_variables)
                prompt_variables = sanitize_variable_values(parsed_vars)
            except (TypeError, ValueError):
                current_app.logger.warning(
                    f"Could not parse prompt_variables form field as JSON: {raw_prompt_variables[:200]}"
                )

        # Convert to int if provided
        if min_speakers:
            try:
                min_speakers = int(min_speakers)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers:
            try:
                max_speakers = int(max_speakers)
            except (ValueError, TypeError):
                max_speakers = None

        # Apply precedence hierarchy: user input > tag defaults > folder defaults > environment variables > user defaults > auto-detect

        # Apply folder defaults first (lower priority than tags)
        if selected_folder and not selected_tags:
            # Only apply folder defaults if no tags are selected (tags take priority)
            if not language and selected_folder.default_language:
                language = selected_folder.default_language
            if min_speakers is None and selected_folder.default_min_speakers:
                min_speakers = selected_folder.default_min_speakers
            if max_speakers is None and selected_folder.default_max_speakers:
                max_speakers = selected_folder.default_max_speakers
            if not hotwords and selected_folder.default_hotwords:
                hotwords = selected_folder.default_hotwords
            if not initial_prompt and selected_folder.default_initial_prompt:
                initial_prompt = selected_folder.default_initial_prompt
            if not transcription_model and selected_folder.default_transcription_model:
                transcription_model = selected_folder.default_transcription_model

        # Apply tag defaults if tags are selected and values are not explicitly provided by user
        # Use first tag's defaults (highest priority - overrides folder)
        if selected_tags:
            first_tag = selected_tags[0]
            if not language and first_tag.default_language:
                language = first_tag.default_language
            if min_speakers is None and first_tag.default_min_speakers:
                min_speakers = first_tag.default_min_speakers
            if max_speakers is None and first_tag.default_max_speakers:
                max_speakers = first_tag.default_max_speakers
            if not hotwords and first_tag.default_hotwords:
                hotwords = first_tag.default_hotwords
            if not initial_prompt and first_tag.default_initial_prompt:
                initial_prompt = first_tag.default_initial_prompt
            if not transcription_model and first_tag.default_transcription_model:
                transcription_model = first_tag.default_transcription_model

        # Apply environment variable defaults if still no values are set
        if min_speakers is None and ASR_MIN_SPEAKERS:
            try:
                min_speakers = int(ASR_MIN_SPEAKERS)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers is None and ASR_MAX_SPEAKERS:
            try:
                max_speakers = int(ASR_MAX_SPEAKERS)
            except (ValueError, TypeError):
                max_speakers = None

        # Fall back to user defaults if still not set
        if not language and current_user.transcription_language:
            language = current_user.transcription_language
            current_app.logger.info(f"Using user's default transcription language: {language}")
        if not hotwords and current_user.transcription_hotwords:
            hotwords = current_user.transcription_hotwords
        if not initial_prompt and current_user.transcription_initial_prompt:
            initial_prompt = current_user.transcription_initial_prompt

        # Create initial database entry
        now = datetime.utcnow()

        # Determine meeting_date: prefer user-provided, then client lastModified, then file metadata, then current time
        meeting_date = None

        # First try user-provided meeting_date (ISO 8601 format)
        if user_meeting_date:
            try:
                parsed = datetime.fromisoformat(user_meeting_date.replace('Z', '+00:00'))
                # Strip timezone to store as naive datetime, consistent with other date sources
                meeting_date = parsed.replace(tzinfo=None)
                current_app.logger.info(f"Using user-provided meeting_date: {meeting_date}")
            except (ValueError, TypeError) as e:
                current_app.logger.warning(f"Could not parse user meeting_date '{user_meeting_date}': {e}")

        # Then try client-provided file lastModified (most reliable for uploads)
        if not meeting_date and file_last_modified:
            try:
                # JavaScript lastModified is in milliseconds since epoch
                timestamp_ms = int(file_last_modified)
                meeting_date = datetime.fromtimestamp(timestamp_ms / 1000)
                current_app.logger.info(f"Using client file lastModified: {meeting_date}")
            except (ValueError, TypeError, OSError) as e:
                current_app.logger.warning(f"Could not parse file_last_modified '{file_last_modified}': {e}")

        # Fall back to file metadata (creation_time, date tags, etc.)
        if not meeting_date:
            meeting_date = get_creation_date(filepath, use_file_mtime=False)
            if meeting_date:
                current_app.logger.info(f"Using file metadata creation date: {meeting_date}")

        # Final fallback to current time
        if not meeting_date:
            meeting_date = now
            current_app.logger.debug("No file date available, using current time")

        recording = Recording(
            audio_path=None,
            original_filename=original_filename,
            title=resolve_upload_title(user_title, original_filename),
            file_size=final_file_size,
            status='PENDING',
            meeting_date=meeting_date,
            user_id=current_user.id,
            mime_type=mime_type,
            audio_duration_seconds=audio_duration_seconds,
            notes=notes,
            folder_id=selected_folder.id if selected_folder else None,
            processing_source='upload',  # Track that this was manually uploaded
            file_hash=file_hash,
            prompt_variables=prompt_variables,
            # Per-upload override: record the effective audio-only flag so
            # the processing task knows whether to drop the video stream
            # regardless of the server's VIDEO_RETENTION setting. True when
            # the user explicitly toggled it or when VIDEO_RETENTION is off.
            keep_audio_only=effective_audio_only,
        )
        db.session.add(recording)
        db.session.flush()  # Assign recording.id without committing transaction yet

        storage_key = storage.build_recording_key(original_filename, recording.id, now=now)
        stored_object = storage.upload_local_file(
            filepath,
            storage_key,
            content_type=mime_type,
            delete_source=True,
        )
        # Intentionally do not reuse `filepath` after upload: with delete_source=True
        # the staging file may already be gone (notably on S3 backend). Workers should
        # always read from recording.audio_path via the storage facade.
        recording.audio_path = stored_object.locator

        # Add tags to recording if selected (preserve order)
        for order, tag in enumerate(selected_tags, 1):
            new_association = RecordingTag(
                recording_id=recording.id,
                tag_id=tag.id,
                order=order,
                added_at=datetime.utcnow()
            )
            db.session.add(new_association)

        db.session.commit()

        if selected_tags:
            tag_names = [tag.name for tag in selected_tags]
            current_app.logger.info(f"Added {len(selected_tags)} tags to recording {recording.id}: {', '.join(tag_names)}")

        current_app.logger.info(f"Initial recording record created with ID: {recording.id}")

        # Validate against admin-curated list and apply admin default when
        # nothing else in the chain set a model.
        transcription_model = _resolve_transcription_model(transcription_model)

        # --- Queue transcription job ---
        first_tag = selected_tags[0] if selected_tags else None
        job_params = {
            'language': language,
            'min_speakers': min_speakers,
            'max_speakers': max_speakers,
            'tag_id': first_tag.id if first_tag else None,
            'hotwords': hotwords,
            'initial_prompt': initial_prompt,
            'transcription_model': transcription_model,
        }

        current_app.logger.info(f"Queueing transcription for recording {recording.id} with params: {job_params}")
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='transcribe',
            params=job_params,
            is_new_upload=True
        )
        current_app.logger.info(f"Transcription job queued for recording ID: {recording.id}")

        # Webhook event (#275). Fan-out happens off-request via the
        # dispatcher; this call only enqueues a delivery row per matching
        # subscription, so it is cheap and safe inside the request path.
        try:
            from src.services.webhook_dispatch import emit_webhook_event
            emit_webhook_event(
                user_id=current_user.id,
                event_type='recording.created',
                data={
                    'recording_id': recording.id,
                    'title': recording.title,
                    'file_size': recording.file_size,
                    'original_filename': recording.original_filename,
                },
            )
        except Exception as e:
            current_app.logger.warning(f"Webhook emit (recording.created) failed: {e}")

        response_data = recording.to_dict(viewer_user=current_user)
        if duplicate_warning:
            response_data['duplicate_warning'] = duplicate_warning
        return jsonify(response_data), 202

    except RequestEntityTooLarge:
        # Report the effective limit that fired (regular vs audio-only)
        # so the frontend can show the right message. effective_limit_mb
        # is set at the top of the function before the size check.
        try:
            limit_mb = float(effective_limit_mb)
        except (NameError, TypeError, ValueError):
            limit_mb = float(current_app.config['MAX_CONTENT_LENGTH']) / (1024 * 1024)
        audio_only_mode = False
        try:
            audio_only_mode = bool(effective_audio_only) and bool(is_likely_video_by_ext)
        except NameError:
            pass
        current_app.logger.warning(
            f"Upload failed: File too large (>{limit_mb}MB, audio_only={audio_only_mode})"
        )
        return jsonify({
            'error': f'File too large. Maximum size is {limit_mb:.0f} MB.',
            'max_size_mb': limit_mb,
            'effective_limit_mb': limit_mb,
            'audio_only_mode': audio_only_mode,
        }), 413
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error during file upload: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred during upload.'}), 500


@recordings_bp.route('/api/recordings/incognito', methods=['POST'])
@login_required
def upload_incognito():
    """
    Process audio in incognito mode - no database storage.
    Returns transcript/summary directly in response.

    This endpoint is designed for HIPAA-friendly transcription where
    audio data is processed but never persisted to the database.
    Results are returned directly and only stored client-side in sessionStorage.
    """
    # Check if incognito mode is enabled
    if not ENABLE_INCOGNITO_MODE:
        return jsonify({'error': 'Incognito mode is not enabled on this server'}), 403

    import tempfile
    from datetime import datetime
    from src.tasks.processing import transcribe_incognito, generate_incognito_summary

    temp_filepath = None

    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        original_filename = file.filename
        safe_filename = secure_filename(original_filename)

        # Get file size
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)

        # Mirror the regular /upload route's option-B size policy:
        # video files (by extension) use the larger video cap because
        # incognito always extracts audio (it never retains the video
        # stream), so it's effectively always in audio-only mode.
        # Audio uploads use the regular cap.
        _VIDEO_EXTS = {'.mp4', '.mov', '.mkv', '.avi', '.webm', '.m4v', '.wmv', '.flv', '.ts', '.mts'}
        _is_video = os.path.splitext(original_filename)[1].lower() in _VIDEO_EXTS
        regular_limit_mb = int(SystemSetting.get_setting('max_file_size_mb', 250))
        audio_only_limit_mb = int(
            SystemSetting.get_setting('max_audio_only_video_size_mb', regular_limit_mb * 4)
        )
        effective_limit_mb = audio_only_limit_mb if _is_video else regular_limit_mb
        if file_size > effective_limit_mb * 1024 * 1024:
            return jsonify({
                'error': f'File too large. Maximum size is {effective_limit_mb} MB.',
                'max_size_mb': float(effective_limit_mb),
            }), 413

        # Save to temp file - use secure temp directory
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'_{safe_filename}') as tmp:
            temp_filepath = tmp.name
            file.save(temp_filepath)
            current_app.logger.info(f"[Incognito] Temp file saved: {temp_filepath}")

        # Get optional parameters
        # Note: Empty string '' means auto-detect, don't convert to None
        language = request.form.get('language', '')
        min_speakers = request.form.get('min_speakers')
        max_speakers = request.form.get('max_speakers')
        auto_summarize = request.form.get('auto_summarize', 'false').lower() == 'true'

        # Convert to int if provided
        if min_speakers:
            try:
                min_speakers = int(min_speakers)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers:
            try:
                max_speakers = int(max_speakers)
            except (ValueError, TypeError):
                max_speakers = None

        # Log only metadata - NEVER log content for HIPAA compliance
        current_app.logger.info(f"[Incognito] Processing request from user {current_user.id}: "
                               f"filename={original_filename}, size={file_size/1024/1024:.2f}MB, "
                               f"language={language}, auto_summarize={auto_summarize}")

        # Perform transcription synchronously (no database operations)
        result = transcribe_incognito(
            filepath=temp_filepath,
            original_filename=original_filename,
            language=language,
            min_speakers=min_speakers,
            max_speakers=max_speakers,
            user=current_user
        )

        if result.get('error'):
            current_app.logger.error(f"[Incognito] Transcription failed: {result['error']}")
            return jsonify({
                'incognito': True,
                'error': result['error']
            }), 500

        # Optionally generate summary
        summary = None
        if auto_summarize and result.get('transcription'):
            current_app.logger.info(f"[Incognito] Auto-summarize requested, generating summary...")
            summary = generate_incognito_summary(result['transcription'], current_user)

        # Build response
        # Render markdown to HTML for summary display
        summary_html = None
        if summary:
            summary_html = md_to_html(summary)

        response_data = {
            'incognito': True,
            'transcription': result.get('transcription'),
            'summary': summary,
            'summary_html': summary_html,
            'title': result.get('title', 'Incognito Recording'),
            'audio_duration_seconds': result.get('audio_duration_seconds'),
            'processing_time_seconds': result.get('processing_time_seconds'),
            'created_at': datetime.utcnow().isoformat() + 'Z',
            'original_filename': original_filename,
            'file_size': file_size
        }

        current_app.logger.info(f"[Incognito] Request completed successfully for user {current_user.id}")

        return jsonify(response_data), 200

    except RequestEntityTooLarge:
        max_size_mb = current_app.config['MAX_CONTENT_LENGTH'] / (1024 * 1024)
        current_app.logger.warning(f"[Incognito] Upload failed: File too large (>{max_size_mb}MB)")
        return jsonify({
            'incognito': True,
            'error': f'File too large. Maximum size is {max_size_mb:.0f} MB.',
            'max_size_mb': max_size_mb
        }), 413

    except Exception as e:
        current_app.logger.error(f"[Incognito] Error during processing: {e}", exc_info=True)
        return jsonify({
            'incognito': True,
            'error': 'An unexpected error occurred during processing.'
        }), 500

    finally:
        # CRITICAL: Always delete temp file for HIPAA compliance
        if temp_filepath and os.path.exists(temp_filepath):
            try:
                os.remove(temp_filepath)
                current_app.logger.info(f"[Incognito] Temp file deleted: {temp_filepath}")
            except Exception as cleanup_error:
                current_app.logger.error(f"[Incognito] Failed to delete temp file {temp_filepath}: {cleanup_error}")


@recordings_bp.route('/api/recordings/incognito/chat', methods=['POST'])
@login_required
def chat_incognito():
    """
    Chat with an incognito recording's transcription.
    Since incognito recordings don't exist in the database, the transcription
    is passed directly in the request.
    """
    # Check if incognito mode is enabled
    if not ENABLE_INCOGNITO_MODE:
        return jsonify({'error': 'Incognito mode is not enabled on this server'}), 403

    from src.tasks.processing import format_transcription_for_llm

    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        transcription = data.get('transcription')
        user_message = data.get('message')
        message_history = data.get('message_history', [])
        participants = data.get('participants', '')
        notes = data.get('notes', '')

        if not transcription:
            return jsonify({'error': 'No transcription provided'}), 400
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400

        # Check if chat client is available
        if chat_client is None:
            return jsonify({'error': 'Chat service is not available (chat client not configured)'}), 503

        # Prepare the system prompt with the transcription
        user_chat_output_language = current_user.output_language if current_user.is_authenticated else None

        language_instruction = ""
        if user_chat_output_language:
            language_instruction = f"Please provide all your responses in {user_chat_output_language}."

        user_name = current_user.name if current_user.is_authenticated and current_user.name else "User"
        user_title = current_user.job_title if current_user.is_authenticated and current_user.job_title else "a professional"
        user_company = current_user.company if current_user.is_authenticated and current_user.company else "their organization"

        formatted_transcription = format_transcription_for_llm(transcription)

        # Get configurable transcript length limit for chat
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            chat_transcript = formatted_transcription
        else:
            chat_transcript = formatted_transcription[:transcript_limit]

        system_prompt = f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {language_instruction} Analyze the following meeting information and respond to the specific request.

Following are the meeting participants and their roles:
{participants or "No specific participants information provided."}

Following is the meeting transcript:
<<start transcript>>
{chat_transcript or "No transcript available."}
<<end transcript>>

Additional context and notes about the meeting:
{notes or "none"}

Note: This is an incognito recording - no data is stored on the server.
"""

        # Prepare messages array with system prompt and conversation history
        messages = [{"role": "system", "content": system_prompt}]
        if message_history:
            messages.extend(message_history)
        messages.append({"role": "user", "content": user_message})

        # Get model info
        chat_model = os.environ.get('TEXT_MODEL_NAME', os.environ.get('OPENAI_CHAT_MODEL', 'gpt-4o-mini'))
        user_id = current_user.id

        current_app.logger.info(f"[Incognito Chat] User {user_id} sending message")

        def generate():
            """Stream the chat response."""
            try:
                response = chat_client.chat.completions.create(
                    model=chat_model,
                    messages=messages,
                    temperature=0.7,
                    stream=True
                )

                for chunk in response:
                    if chunk.choices and len(chunk.choices) > 0:
                        delta = chunk.choices[0].delta
                        if hasattr(delta, 'content') and delta.content:
                            yield f"data: {json.dumps({'content': delta.content})}\n\n"

                yield "data: [DONE]\n\n"

            except Exception as e:
                current_app.logger.error(f"[Incognito Chat] Error during streaming: {e}")
                # Provide more helpful error message for connection issues
                error_msg = str(e).lower()
                if 'connection' in error_msg or 'connect' in error_msg or 'refused' in error_msg:
                    yield f"data: {json.dumps({'error': 'Could not connect to LLM server. Please check that your LLM service is running.'})}\n\n"
                else:
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        current_app.logger.error(f"[Incognito Chat] Error: {e}", exc_info=True)
        # Provide more helpful error message for connection issues
        error_msg = str(e).lower()
        if 'connection' in error_msg or 'connect' in error_msg or 'refused' in error_msg:
            return jsonify({'error': 'Could not connect to LLM server. Please check that your LLM service is running.'}), 503
        return jsonify({'error': 'An error occurred during chat'}), 500


@recordings_bp.route('/api/recordings/incognito/summary', methods=['POST'])
@login_required
def generate_incognito_summary_endpoint():
    """
    Generate summary for an incognito recording on demand.
    Since incognito recordings don't exist in the database, the transcription
    is passed directly in the request.
    """
    # Check if incognito mode is enabled
    if not ENABLE_INCOGNITO_MODE:
        return jsonify({'error': 'Incognito mode is not enabled on this server'}), 403

    from src.tasks.processing import generate_incognito_summary
    from src.utils import md_to_html

    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        transcription = data.get('transcription')
        if not transcription:
            return jsonify({'error': 'No transcription provided'}), 400

        # Check if LLM client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (LLM client not configured)'}), 503

        current_app.logger.info(f"[Incognito Summary] User {current_user.id} requesting summary generation")

        # Generate summary using existing function
        summary = generate_incognito_summary(transcription, current_user)

        if summary:
            summary_html = md_to_html(summary)
            return jsonify({
                'summary': summary,
                'summary_html': summary_html
            })
        else:
            return jsonify({'error': 'Failed to generate summary. Please check that your LLM service is running.'}), 503

    except Exception as e:
        current_app.logger.error(f"[Incognito Summary] Error: {e}", exc_info=True)
        # Provide more helpful error message for connection issues
        error_msg = str(e).lower()
        if 'connection' in error_msg or 'connect' in error_msg or 'refused' in error_msg:
            return jsonify({'error': 'Could not connect to LLM server. Please check that your LLM service is running.'}), 503
        return jsonify({'error': f'Failed to generate summary: {str(e)}'}), 500


# Status Endpoint


@recordings_bp.route('/recording/<int:recording_id>', methods=['DELETE'])
@login_required
def delete_recording(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to delete this recording'}), 403

        # Check deletion permissions (admin-only if USERS_CAN_DELETE is false)
        if not USERS_CAN_DELETE and not current_user.is_admin:
            return jsonify({'error': 'Only administrators can delete recordings'}), 403

        # Delete the audio file first
        try:
            if recording.audio_path:
                storage = get_storage_service()
                storage.delete(recording.audio_path, missing_ok=True)
                current_app.logger.info(f"Deleted audio file via storage backend: {recording.audio_path}")
        except Exception as e:
            current_app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")

        # Log embeddings cleanup for Inquire Mode if enabled
        if ENABLE_INQUIRE_MODE:
            chunk_count = TranscriptChunk.query.filter_by(recording_id=recording_id).count()
            if chunk_count > 0:
                current_app.logger.info(f"Deleting {chunk_count} transcript chunks with embeddings for recording {recording_id}")

        # Delete associated records with NOT NULL recording_id constraints
        from src.models.speaker_snippet import SpeakerSnippet
        deleted_snippets = SpeakerSnippet.query.filter_by(recording_id=recording_id).delete()
        if deleted_snippets > 0:
            current_app.logger.info(f"Deleted {deleted_snippets} speaker snippets for recording {recording_id}")

        from src.models.processing_job import ProcessingJob
        deleted_jobs = ProcessingJob.query.filter_by(recording_id=recording_id).delete()
        if deleted_jobs > 0:
            current_app.logger.info(f"Deleted {deleted_jobs} processing jobs for recording {recording_id}")

        # Capture identity for the webhook event before deletion drops the row.
        _deleted_recording_id = recording.id
        _deleted_recording_title = recording.title
        _deleted_user_id = recording.user_id

        # Delete the database record (cascade will handle chunks/embeddings)
        db.session.delete(recording)
        db.session.commit()
        current_app.logger.info(f"Deleted recording record ID: {recording_id}")

        # Webhook event (#275)
        try:
            from src.services.webhook_dispatch import emit_webhook_event
            emit_webhook_event(
                user_id=_deleted_user_id,
                event_type='recording.deleted',
                data={
                    'recording_id': _deleted_recording_id,
                    'title': _deleted_recording_title,
                },
            )
        except Exception as e:
            current_app.logger.warning(f"Webhook emit (recording.deleted) failed: {e}")

        if ENABLE_INQUIRE_MODE and chunk_count > 0:
            current_app.logger.info(f"Successfully deleted embeddings and chunks for recording {recording_id}")

        # Mark the export file as deleted
        mark_export_as_deleted(recording_id)

        # Clean up orphaned speakers (run after successful deletion)
        # This is a best-effort cleanup; failures are logged but don't affect the delete operation
        try:
            from src.services.speaker_cleanup import cleanup_orphaned_speakers
            speaker_stats = cleanup_orphaned_speakers()
            if speaker_stats.get('speakers_deleted', 0) > 0:
                current_app.logger.info(
                    f"Cleaned up {speaker_stats['speakers_deleted']} orphaned speakers after recording deletion"
                )
        except Exception as cleanup_error:
            # Log the error but don't fail the deletion
            current_app.logger.warning(f"Speaker cleanup after recording deletion failed: {cleanup_error}")

        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while deleting.'}), 500


# --- Inbox and Archive Endpoints ---

@recordings_bp.route('/api/inbox_recordings', methods=['GET'])
@login_required
def get_inbox_recordings():
    """Get recordings that are in the inbox and currently processing."""
    from sqlalchemy import select
    try:
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.is_inbox == True,
            Recording.status.in_(['PENDING', 'PROCESSING', 'SUMMARIZING'])
        ).order_by(Recording.created_at.desc())

        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_list_dict(viewer_user=current_user) for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching inbox recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recordings/archived', methods=['GET'])
@login_required
def get_archived_recordings():
    """Get recordings where audio has been deleted but transcription remains."""
    from sqlalchemy import select
    try:
        search_query = request.args.get('q', '').strip()

        # Find recordings owned by current user where audio_deleted_at is not null
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.audio_deleted_at.is_not(None)
        ).order_by(Recording.audio_deleted_at.desc())

        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_list_dict(viewer_user=current_user) for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching archived recordings: {e}")
        return jsonify({'error': str(e)}), 500


# --- Recording Detail and Audio Endpoints ---

@recordings_bp.route('/api/recordings/<int:recording_id>', methods=['GET'])
@login_required
def get_recording_detail(recording_id):
    """Get full details for a specific recording including markdown HTML."""
    try:
        recording = db.session.get(Recording, recording_id)

        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check ownership or shared access
        has_access = recording.user_id == current_user.id

        # Check if recording has been shared with current user (if internal sharing is enabled)
        if not has_access and ENABLE_INTERNAL_SHARING:
            share = InternalShare.query.filter_by(
                recording_id=recording_id,
                shared_with_user_id=current_user.id
            ).first()
            has_access = share is not None

        if not has_access:
            return jsonify({'error': 'Access denied'}), 403

        # Return full detail with HTML conversion
        rec_dict = recording.to_dict(include_html=True, viewer_user=current_user)

        # Add sharing metadata
        is_owner = recording.user_id == current_user.id
        rec_dict['is_owner'] = is_owner

        # Add edit permission info (uses has_recording_access which checks group admin status)
        rec_dict['can_edit'] = has_recording_access(recording, current_user, require_edit=True)

        # Add delete permission info (only owner can delete)
        rec_dict['can_delete'] = is_owner and (USERS_CAN_DELETE or current_user.is_admin)

        # Add sharing-related fields
        if not is_owner:
            # This is a shared recording - get owner info and share permissions
            owner = db.session.get(User, recording.user_id)
            rec_dict['owner_username'] = owner.username if owner else "Unknown"
            rec_dict['is_shared'] = True
            # Don't show outgoing share count for recordings you don't own
            rec_dict['shared_with_count'] = 0
            rec_dict['public_share_count'] = 0

            # Get share permissions
            share = InternalShare.query.filter_by(
                recording_id=recording.id,
                shared_with_user_id=current_user.id
            ).first()

            if share:
                rec_dict['share_info'] = {
                    'share_id': share.id,
                    'owner_username': owner.username if owner else "Unknown",
                    'can_edit': share.can_edit,
                    'can_reshare': share.can_reshare,
                    'shared_at': share.created_at.isoformat()
                }
            else:
                # Fallback if share record not found (shouldn't happen)
                rec_dict['share_info'] = {
                    'can_edit': False,
                    'can_reshare': False
                }
        else:
            rec_dict['is_shared'] = False

        # Check if recording has group tags (among visible tags)
        visible_tags = recording.get_visible_tags(current_user)
        has_group_tags = any(tag.is_group_tag for tag in visible_tags) if visible_tags else False
        rec_dict['has_group_tags'] = has_group_tags

        # Enrich with per-user status
        enrich_recording_dict_with_user_status(rec_dict, recording, current_user)

        return jsonify(rec_dict)
    except Exception as e:
        current_app.logger.error(f"Error fetching recording detail: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/recording/<int:recording_id>/status', methods=['GET'])
@login_required
def get_recording_status_only(recording_id):
    """
    Lightweight endpoint that returns only the status field.
    Used for polling during processing/summarization.
    Note: Rate limiting exemption is configured at app level.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Return only the status field
        return jsonify({'status': recording.status})
    except Exception as e:
        current_app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/batch-status', methods=['POST'])
@login_required
def get_batch_recording_status():
    """
    Batch endpoint to get status for multiple recordings at once.
    More efficient than polling individual status endpoints.

    Request body: {"recording_ids": [1, 2, 3]}
    Response: {"statuses": {"1": "COMPLETED", "2": "PROCESSING", "3": "FAILED"}}
    """
    try:
        data = request.get_json()
        if not data or 'recording_ids' not in data:
            return jsonify({'error': 'recording_ids is required'}), 400

        recording_ids = data['recording_ids']
        if not isinstance(recording_ids, list):
            return jsonify({'error': 'recording_ids must be a list'}), 400

        # Limit batch size to prevent abuse
        if len(recording_ids) > 50:
            return jsonify({'error': 'Maximum 50 recordings per batch'}), 400

        # Query all recordings at once
        recordings = Recording.query.filter(Recording.id.in_(recording_ids)).all()

        # Build response with only accessible recordings
        statuses = {}
        for recording in recordings:
            if has_recording_access(recording, current_user):
                statuses[str(recording.id)] = recording.status

        return jsonify({'statuses': statuses})
    except Exception as e:
        current_app.logger.error(f"Error fetching batch status: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/job-queue-status', methods=['GET'])
@login_required
def get_job_queue_status():
    """
    Get detailed job queue status for all jobs (active, completed, and failed).
    Returns status for the user's jobs within the last hour.
    """
    try:
        from src.models import ProcessingJob
        from src.services.job_queue import TRANSCRIPTION_JOBS, SUMMARY_JOBS
        from datetime import timedelta

        # Expire all cached objects to ensure we see latest data from worker threads
        db.session.expire_all()

        # Get all jobs for the user (active + recent completed/failed within last hour)
        cutoff_time = datetime.utcnow() - timedelta(hours=1)
        all_jobs = ProcessingJob.query.filter(
            ProcessingJob.user_id == current_user.id,
            db.or_(
                ProcessingJob.status.in_(['queued', 'processing']),
                db.and_(
                    ProcessingJob.status.in_(['completed', 'failed']),
                    ProcessingJob.completed_at >= cutoff_time
                )
            )
        ).order_by(ProcessingJob.created_at.desc()).all()

        job_details = []
        for job in all_jobs:
            recording = db.session.get(Recording, job.recording_id)
            recording_title = None
            if recording:
                recording_title = recording.title or recording.original_filename or 'Untitled'

            # Determine queue type
            queue_type = 'summary' if job.job_type in SUMMARY_JOBS else 'transcription'

            # Calculate position if queued
            position = None
            if job.status == 'queued':
                job_types = SUMMARY_JOBS if job.job_type in SUMMARY_JOBS else TRANSCRIPTION_JOBS
                ahead_in_queue = ProcessingJob.query.filter(
                    ProcessingJob.status == 'queued',
                    ProcessingJob.job_type.in_(job_types),
                    ProcessingJob.created_at < job.created_at
                ).count()
                currently_processing = ProcessingJob.query.filter(
                    ProcessingJob.status == 'processing',
                    ProcessingJob.job_type.in_(job_types)
                ).count()
                position = ahead_in_queue + currently_processing + 1

            job_details.append({
                'id': job.id,
                'recording_id': job.recording_id,
                'recording_title': recording_title,
                'job_status': job.status,
                'job_type': job.job_type,
                'queue_type': queue_type,
                'position': position,
                'is_new_upload': job.is_new_upload,
                'error_message': job.error_message,
                'created_at': job.created_at.isoformat() if job.created_at else None,
                'started_at': job.started_at.isoformat() if job.started_at else None,
                'completed_at': job.completed_at.isoformat() if job.completed_at else None
            })

        return jsonify({'jobs': job_details})
    except Exception as e:
        current_app.logger.error(f"Error fetching job queue status: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/<int:job_id>/retry', methods=['POST'])
@login_required
def retry_failed_job(job_id):
    """Retry a failed job."""
    try:
        from src.models import ProcessingJob

        job = db.session.get(ProcessingJob, job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404

        if job.user_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403

        if job.status != 'failed':
            return jsonify({'error': 'Only failed jobs can be retried'}), 400

        # Reset job for retry
        job.status = 'queued'
        job.error_message = None
        job.retry_count = 0
        job.started_at = None
        job.completed_at = None
        db.session.commit()

        current_app.logger.info(f"Job {job_id} queued for retry by user {current_user.id}")
        return jsonify({'success': True, 'message': 'Job queued for retry'})

    except Exception as e:
        current_app.logger.error(f"Error retrying job {job_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/<int:job_id>', methods=['DELETE'])
@login_required
def delete_job(job_id):
    """Delete a job (clear from queue or history)."""
    try:
        from src.models import ProcessingJob
        import os

        job = db.session.get(ProcessingJob, job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404

        if job.user_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403

        # If it's a failed new upload, also delete the recording
        if job.status == 'failed' and job.is_new_upload:
            recording = db.session.get(Recording, job.recording_id)
            if recording:
                # Delete audio file
                if recording.audio_path:
                    try:
                        get_storage_service().delete(recording.audio_path, missing_ok=True)
                    except Exception as e:
                        current_app.logger.error(f"Error deleting audio file: {e}")
                # Delete ALL processing jobs for this recording first
                ProcessingJob.query.filter_by(recording_id=recording.id).delete()
                db.session.delete(recording)
        else:
            # Just delete this job
            db.session.delete(job)
        db.session.commit()

        current_app.logger.info(f"Job {job_id} deleted by user {current_user.id}")
        return jsonify({'success': True, 'message': 'Job deleted'})

    except Exception as e:
        current_app.logger.error(f"Error deleting job {job_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/clear-completed', methods=['POST'])
@login_required
def clear_completed_jobs():
    """Clear all completed jobs for the current user."""
    try:
        from src.models import ProcessingJob

        deleted = ProcessingJob.query.filter(
            ProcessingJob.user_id == current_user.id,
            ProcessingJob.status == 'completed'
        ).delete(synchronize_session=False)

        db.session.commit()
        current_app.logger.info(f"Cleared {deleted} completed jobs for user {current_user.id}")
        return jsonify({'success': True, 'deleted': deleted})

    except Exception as e:
        current_app.logger.error(f"Error clearing completed jobs: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/status/<int:recording_id>', methods=['GET'])
@login_required
def get_status(recording_id):
    """Endpoint to check the transcription/summarization status (full recording data)."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Ensure events are loaded (refresh the recording to get latest relationships)
        db.session.refresh(recording)

        # Get recording dict and enrich with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)

        return jsonify(recording_dict)
    except Exception as e:
        current_app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/audio/<int:recording_id>')
@login_required
def get_audio(recording_id):
    """Serve audio file for a recording.

    Query parameters:
        download: If 'true', serves file as attachment for download
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording or not recording.audio_path:
            return jsonify({'error': 'Recording or audio file not found'}), 404

        # Check if the recording belongs to the current user or has been shared with them
        has_access = recording.user_id == current_user.id

        # Check if recording has been shared with current user (if internal sharing is enabled)
        if not has_access and ENABLE_INTERNAL_SHARING:
            share = InternalShare.query.filter_by(
                recording_id=recording_id,
                shared_with_user_id=current_user.id
            ).first()
            has_access = share is not None

        if not has_access:
            return jsonify({'error': 'You do not have permission to access this audio file'}), 403
        storage = get_storage_service()

        # Check if download is requested
        download = request.args.get('download', 'false').lower() == 'true'
        download_name = None
        if download:
            # Generate filename from recording title or use default
            filename = recording.title or f'recording_{recording_id}'
            # Sanitize filename and add extension
            filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
            ext = os.path.splitext(recording.original_filename or '')[1] or '.mp3'
            download_name = f"{filename}{ext}"

        delivery = storage.get_audio_delivery(
            recording.audio_path,
            download=download,
            mime_type=recording.mime_type,
            download_name=download_name,
            is_public=False,
        )

        if delivery.mode == 'redirect_url':
            return redirect(delivery.url, code=302)

        if not delivery.local_path or not os.path.exists(delivery.local_path):
            current_app.logger.error(f"Audio file missing from server: {recording.audio_path}")
            return jsonify({'error': 'Audio file missing from server'}), 404

        if download:
            return send_file(
                delivery.local_path,
                as_attachment=True,
                download_name=download_name,
                mimetype=recording.mime_type,
                conditional=True,
            )

        return send_file(delivery.local_path, mimetype=recording.mime_type, conditional=True)
    except Exception as e:
        current_app.logger.error(f"Error serving audio for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# --- Chat with Transcription ---

@recordings_bp.route('/chat', methods=['POST'])
@login_required
def chat_with_transcription():
    """Chat with a specific recording's transcription."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        recording_id = data.get('recording_id')
        user_message = data.get('message')
        message_history = data.get('message_history', [])

        if not recording_id:
            return jsonify({'error': 'No recording ID provided'}), 400
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400

        # Get the recording
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to chat with this recording'}), 403

        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No transcription available for this recording'}), 400

        # Check if transcription is an error message (not actual content)
        if is_transcription_error(recording.transcription):
            return jsonify({'error': 'Cannot chat: transcription failed. Please reprocess the transcription first.'}), 400

        # Check if chat client is available
        if chat_client is None:
            return jsonify({'error': 'Chat service is not available (chat client not configured)'}), 503

        # Prepare the system prompt with the transcription
        user_chat_output_language = current_user.output_language if current_user.is_authenticated else None

        language_instruction = ""
        if user_chat_output_language:
            language_instruction = f"Please provide all your responses in {user_chat_output_language}."

        user_name = current_user.name if current_user.is_authenticated and current_user.name else "User"
        user_title = current_user.job_title if current_user.is_authenticated and current_user.job_title else "a professional"
        user_company = current_user.company if current_user.is_authenticated and current_user.company else "their organization"

        formatted_transcription = format_transcription_for_llm(recording.transcription)

        # Get configurable transcript length limit for chat
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            # No limit
            chat_transcript = formatted_transcription
        else:
            chat_transcript = formatted_transcription[:transcript_limit]

        system_prompt = f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {language_instruction} Analyze the following meeting information and respond to the specific request.

Following are the meeting participants and their roles:
{recording.participants or "No specific participants information provided."}

Following is the meeting transcript:
<<start transcript>>
{chat_transcript or "No transcript available."}
<<end transcript>>

Additional context and notes about the meeting:
{recording.notes or "none"}
"""

        # Prepare messages array with system prompt and conversation history
        messages = [{"role": "system", "content": system_prompt}]
        if message_history:
            messages.extend(message_history)
        messages.append({"role": "user", "content": user_message})

        # Capture context before generator starts (app context may not be available inside generator)
        user_id = current_user.id
        app = current_app._get_current_object()

        def generate():
            # Push app context for entire generator execution
            # This is needed because call_chat_completion checks budget which requires db access
            ctx = app.app_context()
            ctx.push()
            try:
                # Enable streaming with user_id for budget enforcement
                stream = call_chat_completion(
                    messages=messages,
                    temperature=0.7,
                    max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                    stream=True,
                    user_id=user_id,
                    operation_type='chat'
                )

                # Use helper function to process streaming with thinking tag support
                for response in process_streaming_with_thinking(stream, user_id=user_id, operation_type='chat', app=app):
                    yield response

            except TokenBudgetExceeded as e:
                app.logger.warning(f"Token budget exceeded for user {user_id}: {e}")
                yield f"data: {json.dumps({'error': str(e), 'budget_exceeded': True})}\n\n"
            except Exception as e:
                app.logger.error(f"Error during chat stream generation: {str(e)}")
                # Provide more helpful error message for connection issues
                error_msg = str(e).lower()
                if 'connection' in error_msg or 'connect' in error_msg or 'refused' in error_msg:
                    yield f"data: {json.dumps({'error': 'Could not connect to LLM server. Please check that your LLM service is running.'})}\n\n"
                else:
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
            finally:
                ctx.pop()

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        current_app.logger.error(f"Error in chat endpoint: {str(e)}")
        error_msg = str(e).lower()
        if 'connection' in error_msg or 'connect' in error_msg or 'refused' in error_msg:
            return jsonify({'error': 'Could not connect to LLM server. Please check that your LLM service is running.'}), 503
        return jsonify({'error': str(e)}), 500


# --- Tag Management for Recordings ---

@recordings_bp.route('/api/recordings/<int:recording_id>/tags', methods=['POST'])
@login_required
def add_tag_to_recording(recording_id):
    """Add a tag to a recording. Triggers auto-share for group tags."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has view access to this recording
        # (Edit permission will be checked for group tags specifically)
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        data = request.get_json()
        tag_id = data.get('tag_id')

        if not tag_id:
            return jsonify({'error': 'Tag ID is required'}), 400

        tag = db.session.get(Tag, tag_id)
        if not tag:
            return jsonify({'error': 'Tag not found'}), 404

        # Check if user has access to this tag and permission to apply it
        if tag.group_id:
            # Group tag - check membership first
            membership = GroupMembership.query.filter_by(
                group_id=tag.group_id,
                user_id=current_user.id
            ).first()
            if not membership:
                return jsonify({'error': 'You do not have access to this tag'}), 403

            # Only file owner or group admin can apply group tags
            if recording.user_id != current_user.id and membership.role != 'admin':
                return jsonify({'error': 'Only recording owner or group admin can apply group tags'}), 403

            # Group tags require edit permission
            if not has_recording_access(recording, current_user, require_edit=True):
                return jsonify({'error': 'You do not have permission to apply group tags to this recording'}), 403
        else:
            # Personal tag - only the tag owner can use it (view access is sufficient)
            if tag.user_id != current_user.id:
                return jsonify({'error': 'You can only apply your own personal tags'}), 403

        # Check if tag is already on the recording
        existing = RecordingTag.query.filter_by(
            recording_id=recording_id,
            tag_id=tag_id
        ).first()

        if existing:
            return jsonify({'error': 'Tag is already on this recording'}), 400

        # Get the next order position
        max_order = db.session.query(db.func.max(RecordingTag.order)).filter_by(
            recording_id=recording_id
        ).scalar() or 0

        # Add the tag
        recording_tag = RecordingTag(
            recording_id=recording_id,
            tag_id=tag_id,
            order=max_order + 1
        )
        db.session.add(recording_tag)

        # If this is a group tag with sharing enabled, automatically share the recording
        # Only auto-share if recording is completed (not during processing)
        if tag.group_id and ENABLE_INTERNAL_SHARING and recording.status == 'COMPLETED' and (tag.auto_share_on_apply or tag.share_with_group_lead):
            # Determine who to share with
            if tag.auto_share_on_apply:
                group_members = GroupMembership.query.filter_by(group_id=tag.group_id).all()
            elif tag.share_with_group_lead:
                group_members = GroupMembership.query.filter_by(group_id=tag.group_id, role='admin').all()
            else:
                group_members = []

            shares_created = 0
            for membership_to_share in group_members:
                # Skip the recording owner
                if membership_to_share.user_id == recording.user_id:
                    continue

                # Check if already shared
                existing_share = InternalShare.query.filter_by(
                    recording_id=recording_id,
                    shared_with_user_id=membership_to_share.user_id
                ).first()

                if not existing_share:
                    # Create internal share with correct permissions
                    # Group admins get edit permission, regular members get read-only
                    share = InternalShare(
                        recording_id=recording_id,
                        owner_id=recording.user_id,
                        shared_with_user_id=membership_to_share.user_id,
                        can_edit=(membership_to_share.role == 'admin'),
                        can_reshare=False,
                        source_type='group_tag',
                        source_tag_id=tag.id
                    )
                    db.session.add(share)

                    # Check if SharedRecordingState already exists (might exist from previous share)
                    existing_state = SharedRecordingState.query.filter_by(
                        recording_id=recording_id,
                        user_id=membership_to_share.user_id
                    ).first()

                    if not existing_state:
                        # Create SharedRecordingState with default values for the recipient
                        state = SharedRecordingState(
                            recording_id=recording_id,
                            user_id=membership_to_share.user_id,
                            is_inbox=True,  # New shares appear in inbox by default
                            is_highlighted=False  # Not favorited by default
                        )
                        db.session.add(state)

                    shares_created += 1
                    current_app.logger.info(f"Auto-shared recording {recording_id} with user {membership_to_share.user_id} (role={membership_to_share.role}) via group tag '{tag.name}'")

            if shares_created > 0:
                current_app.logger.info(f"Created {shares_created} auto-shares for recording {recording_id} via group tag '{tag.name}'")

        db.session.commit()

        # Return updated recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'recording': recording_dict,
            'tag': tag.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error adding tag to recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/<int:recording_id>/tags/<int:tag_id>', methods=['DELETE'])
@login_required
def remove_tag_from_recording(recording_id, tag_id):
    """Remove a tag from a recording. Cleans up auto-shares for group tags."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has view access to this recording
        # (Edit permission will be checked for group tags specifically)
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Find the recording-tag association
        recording_tag = RecordingTag.query.filter_by(
            recording_id=recording_id,
            tag_id=tag_id
        ).first()

        if not recording_tag:
            return jsonify({'error': 'Tag is not on this recording'}), 404

        # Get the tag to check permissions and for cleanup
        tag = db.session.get(Tag, tag_id)
        if tag:
            # Check permissions to remove this specific tag
            if tag.group_id:
                # Group tag - only file owner or group admin can remove
                membership = GroupMembership.query.filter_by(
                    group_id=tag.group_id,
                    user_id=current_user.id
                ).first()
                if recording.user_id != current_user.id:
                    if not membership or membership.role != 'admin':
                        return jsonify({'error': 'Only recording owner or group admin can remove group tags'}), 403

                # Group tags require edit permission
                if not has_recording_access(recording, current_user, require_edit=True):
                    return jsonify({'error': 'You do not have permission to remove group tags from this recording'}), 403
            else:
                # Personal tag - can be removed by tag owner (view access) or recording owner (edit access)
                if tag.user_id != current_user.id:
                    # Not the tag owner, must be recording owner with edit permission
                    if not has_recording_access(recording, current_user, require_edit=True):
                        return jsonify({'error': 'You can only remove your own personal tags'}), 403

        # Remove the association
        db.session.delete(recording_tag)

        # Clean up shares created by this group tag
        if tag and tag.group_id and ENABLE_INTERNAL_SHARING:
            shares_to_check = InternalShare.query.filter_by(
                recording_id=recording_id,
                source_tag_id=tag_id
            ).all()

            shares_removed = 0
            for share in shares_to_check:
                # Check if user still has access via another group tag on this recording
                other_team_tag_access = db.session.query(Tag).join(
                    RecordingTag, RecordingTag.tag_id == Tag.id
                ).join(
                    GroupMembership, GroupMembership.group_id == Tag.group_id
                ).filter(
                    RecordingTag.recording_id == recording_id,
                    GroupMembership.user_id == share.shared_with_user_id,
                    Tag.id != tag_id,  # Exclude the tag being removed
                    Tag.group_id.isnot(None),
                    db.or_(Tag.auto_share_on_apply == True, Tag.share_with_group_lead == True)
                ).first()

                # Only remove share if user has no other group tag access
                if not other_team_tag_access:
                    db.session.delete(share)
                    shares_removed += 1
                    current_app.logger.info(f"Removed auto-share for user {share.shared_with_user_id} from recording {recording_id} (group tag '{tag.name}' removed)")

            if shares_removed > 0:
                current_app.logger.info(f"Cleaned up {shares_removed} auto-shares for recording {recording_id} after removing group tag '{tag.name}'")

        db.session.commit()

        # Return updated recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'recording': recording_dict
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error removing tag from recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/<int:recording_id>/tags/reorder', methods=['PUT'])
@login_required
def reorder_recording_tags(recording_id):
    """Reorder tags on a recording. Updates the order field for each RecordingTag."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has edit access to this recording
        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        data = request.get_json()
        if not data or 'tag_ids' not in data:
            return jsonify({'error': 'Missing tag_ids in request body'}), 400

        tag_ids = data.get('tag_ids', [])
        if not isinstance(tag_ids, list):
            return jsonify({'error': 'tag_ids must be a list'}), 400

        # Update order for each tag
        for order, tag_id in enumerate(tag_ids, 1):
            recording_tag = RecordingTag.query.filter_by(
                recording_id=recording_id,
                tag_id=tag_id
            ).first()
            if recording_tag:
                recording_tag.order = order

        db.session.commit()

        # Return updated recording
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'recording': recording_dict
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reordering tags on recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# --- Bulk Operations ---

@recordings_bp.route('/api/recordings/bulk', methods=['DELETE'])
@login_required
def bulk_delete_recordings():
    """Delete multiple recordings at once."""
    try:
        data = request.get_json()
        if not data or 'recording_ids' not in data:
            return jsonify({'error': 'Missing recording_ids'}), 400

        recording_ids = data.get('recording_ids', [])
        if not isinstance(recording_ids, list) or len(recording_ids) == 0:
            return jsonify({'error': 'recording_ids must be a non-empty list'}), 400

        # Limit bulk operations to prevent abuse
        if len(recording_ids) > 100:
            return jsonify({'error': 'Cannot delete more than 100 recordings at once'}), 400

        # Check deletion permissions
        if not USERS_CAN_DELETE and not current_user.is_admin:
            return jsonify({'error': 'Only administrators can delete recordings'}), 403

        deleted_ids = []
        errors = []

        for recording_id in recording_ids:
            try:
                recording = db.session.get(Recording, recording_id)
                if not recording:
                    errors.append(f"Recording {recording_id} not found")
                    continue

                # Check ownership
                if recording.user_id and recording.user_id != current_user.id:
                    errors.append(f"No permission for recording {recording_id}")
                    continue

                # Delete audio file
                if recording.audio_path:
                    try:
                        get_storage_service().delete(recording.audio_path, missing_ok=True)
                    except Exception as e:
                        current_app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")

                # Delete associated records with NOT NULL recording_id constraints
                from src.models import ProcessingJob
                from src.models.speaker_snippet import SpeakerSnippet
                SpeakerSnippet.query.filter_by(recording_id=recording_id).delete()
                ProcessingJob.query.filter_by(recording_id=recording_id).delete()

                # Delete the recording
                db.session.delete(recording)
                deleted_ids.append(recording_id)

            except Exception as e:
                current_app.logger.error(f"Error deleting recording {recording_id}: {e}")
                errors.append(f"Error with recording {recording_id}")

        db.session.commit()

        return jsonify({
            'success': True,
            'deleted_ids': deleted_ids,
            'deleted_count': len(deleted_ids),
            'errors': errors if errors else None
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in bulk delete: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred'}), 500


@recordings_bp.route('/api/recordings/bulk-tags', methods=['POST'])
@login_required
def bulk_update_tags():
    """Add or remove a tag from multiple recordings."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Missing request body'}), 400

        recording_ids = data.get('recording_ids', [])
        tag_id = data.get('tag_id')
        action = data.get('action', 'add')  # 'add' or 'remove'

        if not recording_ids or not tag_id:
            return jsonify({'error': 'Missing recording_ids or tag_id'}), 400

        if action not in ['add', 'remove']:
            return jsonify({'error': 'Action must be "add" or "remove"'}), 400

        if len(recording_ids) > 100:
            return jsonify({'error': 'Cannot update more than 100 recordings at once'}), 400

        # Verify tag exists and user has access
        tag = db.session.get(Tag, tag_id)
        if not tag:
            return jsonify({'error': 'Tag not found'}), 404

        if tag.user_id != current_user.id and not tag.group_id:
            return jsonify({'error': 'No permission to use this tag'}), 403

        affected_ids = []

        for recording_id in recording_ids:
            try:
                recording = db.session.get(Recording, recording_id)
                if not recording:
                    continue

                # Check ownership or edit access
                if not has_recording_access(recording, current_user, require_edit=True):
                    continue

                if action == 'add':
                    # Check if tag already exists
                    existing = RecordingTag.query.filter_by(
                        recording_id=recording_id,
                        tag_id=tag_id
                    ).first()

                    if not existing:
                        # Get max order for this recording
                        max_order = db.session.query(db.func.max(RecordingTag.order)).filter_by(
                            recording_id=recording_id
                        ).scalar() or 0

                        new_tag = RecordingTag(
                            recording_id=recording_id,
                            tag_id=tag_id,
                            order=max_order + 1
                        )
                        db.session.add(new_tag)
                        affected_ids.append(recording_id)

                else:  # remove
                    recording_tag = RecordingTag.query.filter_by(
                        recording_id=recording_id,
                        tag_id=tag_id
                    ).first()

                    if recording_tag:
                        db.session.delete(recording_tag)
                        affected_ids.append(recording_id)

            except Exception as e:
                current_app.logger.error(f"Error updating tag for recording {recording_id}: {e}")

        db.session.commit()

        return jsonify({
            'success': True,
            'affected_ids': affected_ids,
            'affected_count': len(affected_ids)
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in bulk tag update: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred'}), 500


@recordings_bp.route('/api/recordings/bulk-reprocess', methods=['POST'])
@login_required
def bulk_reprocess():
    """Queue multiple recordings for reprocessing."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Missing request body'}), 400

        recording_ids = data.get('recording_ids', [])
        reprocess_type = data.get('type', 'summary')  # 'transcription' or 'summary'

        if not recording_ids:
            return jsonify({'error': 'Missing recording_ids'}), 400

        if reprocess_type not in ['transcription', 'summary']:
            return jsonify({'error': 'Type must be "transcription" or "summary"'}), 400

        if len(recording_ids) > 50:
            return jsonify({'error': 'Cannot reprocess more than 50 recordings at once'}), 400

        queued_ids = []

        for recording_id in recording_ids:
            try:
                recording = db.session.get(Recording, recording_id)
                if not recording:
                    continue

                # Check ownership
                if recording.user_id != current_user.id:
                    continue

                # Only reprocess completed or failed recordings
                if recording.status not in ['COMPLETED', 'FAILED']:
                    continue

                # For transcription reprocess, need audio file
                if reprocess_type == 'transcription':
                    if not recording.audio_path or not get_storage_service().exists(recording.audio_path):
                        continue
                    job_type = 'reprocess_transcription'
                else:
                    # For summary, need transcription
                    if not recording.transcription:
                        continue
                    job_type = 'reprocess_summary'

                # Queue the job
                job_queue.enqueue(
                    user_id=current_user.id,
                    recording_id=recording.id,
                    job_type=job_type,
                    params={'user_id': current_user.id}
                )

                queued_ids.append(recording_id)

            except Exception as e:
                current_app.logger.error(f"Error queueing reprocess for recording {recording_id}: {e}")

        db.session.commit()

        return jsonify({
            'success': True,
            'queued_ids': queued_ids,
            'queued_count': len(queued_ids)
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in bulk reprocess: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred'}), 500


@recordings_bp.route('/api/recordings/bulk-toggle', methods=['POST'])
@login_required
def bulk_toggle():
    """Toggle inbox or highlight for multiple recordings."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Missing request body'}), 400

        recording_ids = data.get('recording_ids', [])
        field = data.get('field')  # 'inbox' or 'highlight'
        value = data.get('value')  # True or False

        if not recording_ids or field is None or value is None:
            return jsonify({'error': 'Missing recording_ids, field, or value'}), 400

        if field not in ['inbox', 'highlight']:
            return jsonify({'error': 'Field must be "inbox" or "highlight"'}), 400

        if len(recording_ids) > 100:
            return jsonify({'error': 'Cannot update more than 100 recordings at once'}), 400

        affected_ids = []

        for recording_id in recording_ids:
            try:
                recording = db.session.get(Recording, recording_id)
                if not recording:
                    continue

                # Use set_user_recording_status which handles both owners and shared users
                if field == 'inbox':
                    set_user_recording_status(recording, current_user, is_inbox=value)
                else:
                    set_user_recording_status(recording, current_user, is_highlighted=value)

                affected_ids.append(recording_id)

            except Exception as e:
                current_app.logger.error(f"Error toggling {field} for recording {recording_id}: {e}")

        db.session.commit()

        return jsonify({
            'success': True,
            'affected_ids': affected_ids,
            'affected_count': len(affected_ids)
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in bulk toggle: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred'}), 500


# --- Auto-deletion and Chunks Processing ---

@recordings_bp.route('/api/recordings/<int:recording_id>/toggle_deletion_exempt', methods=['POST'])
@login_required
def toggle_deletion_exempt(recording_id):
    """Toggle the deletion_exempt flag for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check ownership
        if recording.user_id != current_user.id and not current_user.is_admin:
            return jsonify({'error': 'Permission denied'}), 403

        # Toggle the flag
        recording.deletion_exempt = not recording.deletion_exempt
        db.session.commit()

        return jsonify({
            'success': True,
            'deletion_exempt': recording.deletion_exempt
        })
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling deletion exempt for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recording/<int:recording_id>/process_chunks', methods=['POST'])
@login_required
def process_recording_chunks_endpoint(recording_id):
    """Process chunks for a specific recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id != current_user.id:
            return jsonify({'error': 'Permission denied'}), 403

        success = process_recording_chunks(recording_id)
        if success:
            return jsonify({'message': 'Chunks processed successfully'})
        else:
            return jsonify({'error': 'Failed to process chunks'}), 500

    except Exception as e:
        current_app.logger.error(f"Error in process chunks endpoint: {e}")
        return jsonify({'error': str(e)}), 500


# --- Inquire Mode API Endpoints ---
